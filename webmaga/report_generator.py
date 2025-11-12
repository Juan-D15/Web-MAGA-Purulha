"""
Módulo para generar reportes en PDF y Word
Utiliza las plantillas PlantillaPDF.pdf y PlantillaWord.docx
"""

import os
from datetime import datetime
from django.conf import settings
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO

# Intentar importar xhtml2pdf (compatible con Windows)
try:
    from xhtml2pdf import pisa
    XHTML2PDF_AVAILABLE = True
except ImportError:
    XHTML2PDF_AVAILABLE = False
    # Intentar weasyprint como alternativa (requiere dependencias del sistema)
    try:
        import weasyprint
        WEASYPRINT_AVAILABLE = True
    except ImportError:
        WEASYPRINT_AVAILABLE = False


# Información del pie de página
FOOTER_EMAIL = "magabvpurulha@gmail.com"
FOOTER_UBICACION = "km 164.5, CA 14"


def get_template_path(template_name):
    """Obtiene la ruta completa de la plantilla"""
    return os.path.join(settings.BASE_DIR, template_name)


def format_date(date_str):
    """Formatea una fecha en formato legible"""
    if not date_str or date_str == '-':
        return '-'
    try:
        if isinstance(date_str, str):
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        else:
            date_obj = date_str
        return date_obj.strftime('%d/%m/%Y')
    except:
        return str(date_str)


def format_table_word(table):
    """Formatea una tabla Word con bordes visibles y centrado"""
    try:
        # Centrar la tabla
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Agregar bordes a todas las celdas
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        # Definir bordes para la tabla
        tbl = table._tbl
        tblBorders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblPr.append(tblBorders)
        
        # Agregar bordes a cada celda
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.tcPr
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)
                
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
        
    except Exception as e:
        print(f"Advertencia al formatear tabla: {str(e)}")
    
    return table


def add_footer_word(doc):
    """Agrega pie de página al documento Word"""
    try:
        section = doc.sections[0]
        footer = section.footer
        
        # Limpiar footer existente - eliminar todos los párrafos
        for paragraph in list(footer.paragraphs):
            p = paragraph._element
            p.getparent().remove(p)
        
        # Crear párrafo para el pie de página
        footer_para = footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar información
        footer_run = footer_para.add_run(f"Correo: {FOOTER_EMAIL} | Ubicación: {FOOTER_UBICACION}")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
    except Exception as e:
        # Si hay error al agregar footer, no fallar completamente
        print(f"Advertencia: No se pudo agregar pie de página a Word: {str(e)}")
    
    return doc


def add_footer_pdf(html_content):
    """Agrega pie de página al HTML para PDF"""
    footer_html = f"""
    <div style="position: fixed; bottom: 0; width: 100%; text-align: center; 
                font-size: 9pt; color: #808080; padding: 10px; border-top: 1px solid #ddd;">
        Correo: {FOOTER_EMAIL} | Ubicación: {FOOTER_UBICACION}
    </div>
    """
    # Insertar antes del cierre del body
    html_content = html_content.replace('</body>', footer_html + '</body>')
    return html_content


def generate_word_report(report_type, report_data, filters_info=None):
    """Genera un reporte en formato Word"""
    try:
        # Cargar plantilla Word si existe
        template_path = get_template_path('PlantillaWord.docx')
        if os.path.exists(template_path):
            try:
                doc = Document(template_path)
                # Limpiar contenido existente pero mantener estilos y formato
                # Eliminar todos los párrafos de la plantilla
                for paragraph in list(doc.paragraphs):
                    p = paragraph._element
                    p.getparent().remove(p)
            except Exception as e:
                # Si hay error al cargar la plantilla, crear documento nuevo
                print(f"Error al cargar plantilla Word: {str(e)}")
                doc = Document()
        else:
            # Si no existe la plantilla, crear documento nuevo
            doc = Document()
        
        # Agregar título (más compacto, sin filtros)
        # Para evento individual, el título se agrega dentro de la función específica
        if report_type != 'reporte-evento-individual':
            title = doc.add_heading(f'{get_report_title(report_type)}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Agregar fecha de generación (más compacto)
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_run = date_para.add_run(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
            date_run.font.size = Pt(9)
            date_run.font.color.rgb = RGBColor(102, 102, 102)
        
        # Generar contenido según el tipo de reporte
        generate_word_content(doc, report_type, report_data)
        
        # Agregar pie de página
        add_footer_word(doc)
        
        # Guardar en BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise Exception(f"Error al generar reporte Word: {str(e)}")


def get_image_base64(image_path):
    """Convierte una imagen a base64 para incluirla en el PDF"""
    try:
        import base64
        full_path = os.path.join(settings.BASE_DIR, 'src', 'static', image_path)
        if os.path.exists(full_path):
            with open(full_path, 'rb') as img_file:
                img_data = img_file.read()
                img_base64 = base64.b64encode(img_data).decode('utf-8')
                # Determinar el tipo de imagen por extensión
                ext = os.path.splitext(full_path)[1].lower()
                mime_type = 'image/png' if ext == '.png' else 'image/jpeg'
                return f'data:{mime_type};base64,{img_base64}'
        return None
    except Exception as e:
        print(f"Error al convertir imagen a base64: {str(e)}")
        return None


def generate_reportlab_content(elements, report_type, report_data, subtitle_style, section_style, normal_style):
    """Genera el contenido específico del reporte usando ReportLab"""
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.units import cm
    
    try:
        if report_type == 'reporte-evento-individual':
            # Reporte de evento individual
            if isinstance(report_data, dict) and 'evento' in report_data:
                evento = report_data['evento']
            elif isinstance(report_data, dict) and 'nombre' in report_data:
                evento = report_data
            else:
                evento = {}
            
            if evento:
                # Información básica del evento
                elements.append(Paragraph("Información del Evento", subtitle_style))
                elements.append(Spacer(1, 0.3*cm))
                
                info_data = [
                    ['Nombre:', evento.get('nombre', '-')],
                    ['Tipo:', evento.get('tipo', '-')],
                    ['Estado:', evento.get('estado', '-')],
                    ['Fecha:', evento.get('fecha', '-')],
                    ['Comunidad:', evento.get('comunidad', '-')],
                    ['Región:', evento.get('region', '-')],
                ]
                
                info_table = Table(info_data, colWidths=[4*cm, 12*cm])
                info_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f7ff')),
                    ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#2c3e50')),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#ddd')),
                ]))
                elements.append(info_table)
                elements.append(Spacer(1, 0.5*cm))
                
                # Descripción
                if evento.get('descripcion'):
                    elements.append(Paragraph("Descripción", section_style))
                    elements.append(Paragraph(evento['descripcion'], normal_style))
                    elements.append(Spacer(1, 0.3*cm))
                
                # Beneficiarios
                if evento.get('beneficiarios'):
                    elements.append(PageBreak())
                    elements.append(Paragraph("Beneficiarios", subtitle_style))
                    elements.append(Spacer(1, 0.3*cm))
                    
                    ben_headers = ['Nombre', 'Tipo', 'Descripción']
                    ben_data = [ben_headers]
                    
                    for ben in evento['beneficiarios'][:50]:  # Limitar a 50
                        ben_data.append([
                            ben.get('nombre', '-'),
                            ben.get('tipo', '-'),
                            ben.get('descripcion', '-')[:50] + '...' if len(ben.get('descripcion', '')) > 50 else ben.get('descripcion', '-')
                        ])
                    
                    ben_table = Table(ben_data, colWidths=[5*cm, 3*cm, 8*cm])
                    ben_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
                    ]))
                    elements.append(ben_table)
                    elements.append(Spacer(1, 0.3*cm))
                
                # Avances
                if evento.get('avances'):
                    elements.append(PageBreak())
                    elements.append(Paragraph("Avances del Evento", subtitle_style))
                    elements.append(Spacer(1, 0.3*cm))
                    
                    for i, avance in enumerate(evento['avances'][:20], 1):  # Limitar a 20
                        elements.append(Paragraph(f"<b>Avance #{i}</b>", section_style))
                        avance_info = f"""
                        <b>Fecha:</b> {avance.get('fecha', '-')}<br/>
                        <b>Descripción:</b> {avance.get('descripcion', '-')}<br/>
                        <b>Realizado por:</b> {avance.get('realizado_por', '-')}
                        """
                        elements.append(Paragraph(avance_info, normal_style))
                        elements.append(Spacer(1, 0.3*cm))
            else:
                elements.append(Paragraph("No se encontró información del evento.", normal_style))
        
        elif report_type in ['actividades-por-region-comunidad', 'beneficiarios-por-region-comunidad', 
                            'actividad-de-personal', 'avances-eventos-generales', 'comunidades', 
                            'actividad-usuarios', 'reporte-general']:
            # Para otros tipos de reportes, generar tablas genéricas
            if isinstance(report_data, list) and len(report_data) > 0:
                # Generar tabla con los datos
                if len(report_data) > 0:
                    # Obtener las claves del primer elemento
                    keys = list(report_data[0].keys())
                    
                    # Crear encabezados
                    headers = [key.replace('_', ' ').title() for key in keys]
                    table_data = [headers]
                    
                    # Agregar datos (limitar a 100 filas)
                    for item in report_data[:100]:
                        row = []
                        for key in keys:
                            value = item.get(key, '-')
                            # Truncar valores largos
                            if isinstance(value, str) and len(value) > 60:
                                value = value[:57] + '...'
                            row.append(str(value))
                        table_data.append(row)
                    
                    # Calcular anchos de columna dinámicamente
                    num_cols = len(keys)
                    col_width = 16.0 / num_cols  # 16cm total dividido por número de columnas
                    col_widths = [col_width * cm] * num_cols
                    
                    table = Table(table_data, colWidths=col_widths)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 8),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                        ('TOPPADDING', (0, 0), (-1, -1), 5),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
                    ]))
                    elements.append(table)
                    elements.append(Spacer(1, 0.5*cm))
                    
                    # Mostrar total de registros
                    total_text = f"<i>Total de registros: {len(report_data)}</i>"
                    if len(report_data) > 100:
                        total_text += f" (mostrando primeros 100)"
                    elements.append(Paragraph(total_text, normal_style))
            elif isinstance(report_data, dict):
                # Si es un diccionario, mostrar pares clave-valor
                for key, value in report_data.items():
                    if isinstance(value, (list, dict)):
                        elements.append(Paragraph(f"<b>{key.replace('_', ' ').title()}:</b>", section_style))
                        elements.append(Paragraph(str(value)[:500], normal_style))
                    else:
                        elements.append(Paragraph(f"<b>{key.replace('_', ' ').title()}:</b> {value}", normal_style))
                    elements.append(Spacer(1, 0.2*cm))
            else:
                elements.append(Paragraph("No hay datos disponibles para este reporte.", normal_style))
        
        else:
            elements.append(Paragraph("Tipo de reporte no reconocido.", normal_style))
    
    except Exception as e:
        print(f"Error generando contenido ReportLab: {e}")
        import traceback
        traceback.print_exc()
        elements.append(Paragraph(f"Error al generar el contenido: {str(e)}", normal_style))


def generate_pdf_report(report_type, report_data, filters_info=None):
    """Genera un reporte en formato PDF usando ReportLab directamente"""
    try:
        # Intentar usar ReportLab para generar PDF nativo (mejor calidad)
        try:
            from reportlab.lib.pagesizes import A4, letter
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch, cm
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
            from reportlab.platypus.flowables import HRFlowable
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
            
            # Crear buffer para el PDF
            pdf_file = BytesIO()
            
            # Crear documento con márgenes
            doc = SimpleDocTemplate(
                pdf_file,
                pagesize=A4,
                rightMargin=2*cm,
                leftMargin=2*cm,
                topMargin=2.5*cm,
                bottomMargin=2.5*cm,
                title=get_report_title(report_type)
            )
            
            # Contenedor para elementos del documento
            elements = []
            
            # Estilos
            styles = getSampleStyleSheet()
            
            # Estilo personalizado para título
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=12,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            # Estilo para subtítulos
            subtitle_style = ParagraphStyle(
                'CustomSubtitle',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#0772d2'),
                spaceAfter=10,
                spaceBefore=16,
                fontName='Helvetica-Bold',
                leftIndent=10,
                borderPadding=8,
                backColor=colors.HexColor('#f0f7ff'),
                borderColor=colors.HexColor('#0772d2'),
                borderWidth=0,
                borderRadius=4
            )
            
            # Estilo para secciones
            section_style = ParagraphStyle(
                'CustomSection',
                parent=styles['Heading3'],
                fontSize=12,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=8,
                spaceBefore=12,
                fontName='Helvetica-Bold'
            )
            
            # Estilo para texto normal
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#333333'),
                spaceAfter=6,
                leading=14,
                fontName='Helvetica'
            )
            
            # Agregar logo si existe
            logo_path = os.path.join(settings.BASE_DIR, 'src', 'static', 'img', 'logos', 'logo maga letras png.png')
            if os.path.exists(logo_path):
                try:
                    logo = Image(logo_path, width=2.5*cm, height=2.5*cm)
                    logo.hAlign = 'CENTER'
                    elements.append(logo)
                    elements.append(Spacer(1, 0.3*cm))
                except Exception as e:
                    print(f"Error cargando logo: {e}")
            
            # Título del reporte
            title = Paragraph(get_report_title(report_type), title_style)
            elements.append(title)
            elements.append(Spacer(1, 0.3*cm))
            
            # Información de filtros
            if filters_info:
                date_info = Paragraph(f"<i>Filtros aplicados: {filters_info}</i>", normal_style)
                elements.append(date_info)
                elements.append(Spacer(1, 0.3*cm))
            
            # Fecha de generación
            fecha_generacion = datetime.now().strftime('%d/%m/%Y %H:%M')
            date_para = Paragraph(f"<i>Generado el: {fecha_generacion}</i>", normal_style)
            elements.append(date_para)
            elements.append(Spacer(1, 0.5*cm))
            
            # Línea separadora
            elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e0e0e0')))
            elements.append(Spacer(1, 0.5*cm))
            
            # Generar contenido específico del reporte usando ReportLab
            generate_reportlab_content(elements, report_type, report_data, subtitle_style, section_style, normal_style)
            
            # Pie de página con información de contacto
            elements.append(Spacer(1, 1*cm))
            elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e0e0e0')))
            elements.append(Spacer(1, 0.3*cm))
            footer_text = f"<i>Correo: {FOOTER_EMAIL} | Ubicación: {FOOTER_UBICACION}</i>"
            footer_para = Paragraph(footer_text, ParagraphStyle('Footer', parent=normal_style, fontSize=8, textColor=colors.grey, alignment=TA_CENTER))
            elements.append(footer_para)
            
            # Construir el PDF
            doc.build(elements)
            pdf_file.seek(0)
            return pdf_file
            
        except ImportError as ie:
            print(f"ReportLab no disponible: {ie}")
            # Fallback al método HTML si ReportLab no está disponible
            pass
        except Exception as e:
            print(f"Error al generar PDF con ReportLab: {str(e)}")
            import traceback
            traceback.print_exc()
            # Fallback al método HTML
            pass
        
        # Fallback: Generar HTML del reporte (método anterior)
        html_content = generate_html_content(report_type, report_data, filters_info)
        
        # Obtener imágenes en base64
        logo_header = get_image_base64('img/logos/logo maga letras png.png')
        logo_watermark = get_image_base64('img/logos/maga_logo.png')
        
        # Agregar estilos CSS (compatible con xhtml2pdf)
        # Usar @page para pie de página y encabezado con texto
        header_html = ""
        if logo_header:
            # Logo más pequeño, similar al tamaño en Word
            header_html = f"""
            <div class="page-header">
                <img src="{logo_header}" style="max-width: 100px; height: auto;" />
            </div>
            """
        
        watermark_style = ""
        if logo_watermark:
            watermark_style = f"""
            background-image: url('{logo_watermark}');
            background-size: 350px 350px;
            background-repeat: no-repeat;
            background-position: center;
            opacity: 0.1;
            """
        
        footer_text = f"Correo: {FOOTER_EMAIL} | Ubicación: {FOOTER_UBICACION}"
        
        css_styles = f"""
        <style>
            @page {{
                size: A4;
                margin: 2.5cm 2cm 2.5cm 2cm;
            }}
            .page-header {{
                text-align: center;
                padding: 0.2cm 0;
                margin-bottom: 0.3cm;
            }}
            .page-footer {{
                position: fixed;
                bottom: 0;
                left: 0;
                right: 0;
                text-align: center;
                font-size: 8pt;
                color: #666;
                padding: 0.5cm;
                background-color: white;
                border-top: 1px solid #ddd;
            }}
            body {{
                font-family: Arial, sans-serif;
                font-size: 10pt;
                line-height: 1.4;
                color: #333;
                margin-bottom: 2cm;
                {watermark_style}
            }}
            h1 {{
                color: #2c3e50;
                text-align: center;
                font-size: 16pt;
                margin-top: 0.2cm;
                margin-bottom: 0.5cm;
                padding: 0;
            }}
            h2 {{
                color: #0772d2;
                font-size: 14pt;
                font-weight: bold;
                margin-top: 0.8cm;
                margin-bottom: 0.5cm;
                padding: 0.3cm 0.5cm;
                background: linear-gradient(135deg, #f0f7ff 0%, #e6f2ff 100%);
                border-left: 4px solid #0772d2;
                border-radius: 4px;
                page-break-before: always;
            }}
            h2:first-of-type {{
                page-break-before: auto !important;
            }}
            .section-break {{
                page-break-before: always;
            }}
            .section-break:first-child {{
                page-break-before: auto !important;
            }}
            h3 {{
                color: #2c3e50;
                font-size: 12pt;
                font-weight: 600;
                margin-top: 0.6cm;
                margin-bottom: 0.4cm;
                padding-bottom: 0.2cm;
                border-bottom: 2px solid #e0e0e0;
            }}
            p {{
                margin: 0.25cm 0;
                line-height: 1.5;
            }}
            .date-info {{
                text-align: right;
                font-size: 9pt;
                color: #666;
                margin-bottom: 0.5cm;
                font-style: italic;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 0.5cm auto;
                font-size: 9pt;
                page-break-inside: avoid;
                border: 1px solid #000;
            }}
            th {{
                background-color: #2c3e50;
                color: white;
                padding: 6px;
                text-align: left;
                font-weight: bold;
                border: 1px solid #000;
            }}
            td {{
                padding: 5px;
                border: 1px solid #000;
            }}
            tr:nth-child(even) {{
                background-color: #f5f5f5;
            }}
            .metric-box {{
                background: linear-gradient(135deg, #ffffff 0%, #f8f9fa 100%);
                padding: 0.5cm 0.6cm;
                margin: 0.5cm 0;
                border: 1px solid #e0e0e0;
                border-left: 4px solid #0772d2;
                border-radius: 6px;
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
            }}
            .metric-label {{
                font-weight: 600;
                color: #2c3e50;
                font-size: 9.5pt;
                display: inline-block;
                min-width: 180px;
            }}
            .metric-value {{
                font-size: 16pt;
                color: #0772d2;
                margin-top: 2px;
                font-weight: bold;
            }}
            ul, ol {{
                margin: 0.3cm 0;
                padding-left: 1.5cm;
            }}
            li {{
                margin: 0.2cm 0;
            }}
        </style>
        """
        
        footer_html = f"""
        <div class="page-footer">
            {footer_text}
        </div>
        """
        
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            {css_styles}
        </head>
        <body>
            {header_html}
            {html_content}
            {footer_html}
        </body>
        </html>
        """
        
        # Generar PDF usando xhtml2pdf (compatible con Windows)
        pdf_file = BytesIO()
        
        if XHTML2PDF_AVAILABLE:
            # Usar xhtml2pdf
            result = pisa.CreatePDF(
                src=full_html,
                dest=pdf_file,
                encoding='utf-8'
            )
            
            if result.err:
                raise Exception(f"Error al generar PDF con xhtml2pdf: {result.err}")
                
        elif WEASYPRINT_AVAILABLE:
            # Usar weasyprint como alternativa
            weasyprint.HTML(string=full_html).write_pdf(pdf_file)
        else:
            raise Exception("No hay librerías de PDF disponibles. Instale xhtml2pdf o weasyprint.")
        
        pdf_file.seek(0)
        return pdf_file
        
    except Exception as e:
        raise Exception(f"Error al generar reporte PDF: {str(e)}")


def get_report_title(report_type):
    """Obtiene el título del reporte según su tipo"""
    titles = {
        'actividades-por-region-comunidad': 'Actividades por Región o Comunidad',
        'beneficiarios-por-region-comunidad': 'Beneficiarios por Región o Comunidad',
        'actividad-de-personal': 'Actividad de Personal',
        'avances-eventos-generales': 'Avances de Eventos Generales',
        'reporte-evento-individual': 'Reporte de Evento Individual',
        'comunidades': 'Reporte de Comunidades',
        'actividad-usuarios': 'Actividad de Usuarios',
        'reporte-general': 'Reporte General'
    }
    return titles.get(report_type, 'Reporte')


def generate_word_content(doc, report_type, report_data):
    """Genera el contenido específico del reporte en Word"""
    
    if report_type == 'actividades-por-region-comunidad':
        generate_word_actividades_region_comunidad(doc, report_data)
    elif report_type == 'beneficiarios-por-region-comunidad':
        generate_word_beneficiarios_region_comunidad(doc, report_data)
    elif report_type == 'actividad-de-personal':
        generate_word_actividad_personal(doc, report_data)
    elif report_type == 'avances-eventos-generales':
        generate_word_avances_eventos(doc, report_data)
    elif report_type == 'reporte-evento-individual':
        generate_word_evento_individual(doc, report_data)
    elif report_type == 'comunidades':
        generate_word_comunidades(doc, report_data)
    elif report_type == 'actividad-usuarios':
        generate_word_actividad_usuarios(doc, report_data)
    elif report_type == 'reporte-general':
        generate_word_reporte_general(doc, report_data)
    else:
        doc.add_paragraph('Tipo de reporte no reconocido.')


def generate_html_content(report_type, report_data, filters_info=None):
    """Genera el contenido HTML del reporte para PDF"""
    
    html_parts = []
    
    # Título (sin líneas azules, más compacto)
    html_parts.append(f'<h1>{get_report_title(report_type)}</h1>')
    
    # Fecha de generación (más compacto)
    html_parts.append(f'<p class="date-info">Fecha de generación: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}</p>')
    
    # Contenido según tipo
    if report_type == 'actividades-por-region-comunidad':
        html_parts.append(generate_html_actividades_region_comunidad(report_data))
    elif report_type == 'beneficiarios-por-region-comunidad':
        html_parts.append(generate_html_beneficiarios_region_comunidad(report_data))
    elif report_type == 'actividad-de-personal':
        html_parts.append(generate_html_actividad_personal(report_data))
    elif report_type == 'avances-eventos-generales':
        html_parts.append(generate_html_avances_eventos(report_data))
    elif report_type == 'reporte-evento-individual':
        html_parts.append(generate_html_evento_individual(report_data))
    elif report_type == 'comunidades':
        html_parts.append(generate_html_comunidades(report_data))
    elif report_type == 'actividad-usuarios':
        html_parts.append(generate_html_actividad_usuarios(report_data))
    elif report_type == 'reporte-general':
        html_parts.append(generate_html_reporte_general(report_data))
    else:
        html_parts.append('<p>Tipo de reporte no reconocido.</p>')
    
    return '\n'.join(html_parts)


# ========== FUNCIONES ESPECÍFICAS PARA WORD ==========

def generate_word_actividades_region_comunidad(doc, data):
    """Genera contenido Word para actividades por región/comunidad"""
    if not data or len(data) == 0:
        doc.add_paragraph('No se encontraron datos para este reporte.')
        return
    
    # Separar items con resumen (sin actividades) de la entrada con todas las actividades
    items_resumen = []
    todas_actividades_item = None
    
    for item in data:
        # Si tiene actividades y es la entrada de "Todas las Actividades" o similar
        if item.get('actividades') and len(item.get('actividades', [])) > 0:
            todas_actividades_item = item
        else:
            items_resumen.append(item)
    
    # ========== SECCIÓN 1: TABLAS DE RESUMEN (2 por página) ==========
    for idx, item in enumerate(items_resumen):
        # Salto de página cada 2 items (página nueva cada 2 tablas)
        if idx > 0 and idx % 2 == 0:
            doc.add_page_break()
        
        # Título de la comunidad/región
        heading = doc.add_heading(item.get('nombre', 'N/A'), level=3)
        
        # Crear tabla de resumen para esta comunidad/región
        table = doc.add_table(rows=1, cols=2)
        format_table_word(table)
        
        # Ajustar ancho de columnas (más espacio para 2 por página)
        from docx.shared import Inches
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(4.5)
        
        # Encabezado de la tabla
        header_cells = table.rows[0].cells
        header_cells[0].text = "Información"
        header_cells[1].text = "Valor"
        header_cells[0].paragraphs[0].runs[0].bold = True
        header_cells[1].paragraphs[0].runs[0].bold = True
        
        # Agregar filas de datos
        row = table.add_row()
        row.cells[0].text = "Región"
        row.cells[1].text = str(item.get('region', '-'))
        
        row = table.add_row()
        row.cells[0].text = "Total de actividades"
        row.cells[1].text = str(item.get('total_actividades', 0))
        
        row = table.add_row()
        row.cells[0].text = "Total de beneficiarios"
        row.cells[1].text = str(item.get('total_beneficiarios', 0))
        
        row = table.add_row()
        row.cells[0].text = "Beneficiarios individuales"
        row.cells[1].text = str(item.get('beneficiarios_individuales', 0))
        
        row = table.add_row()
        row.cells[0].text = "Beneficiarios familias"
        row.cells[1].text = str(item.get('beneficiarios_familias', 0))
        
        row = table.add_row()
        row.cells[0].text = "Beneficiarios instituciones"
        row.cells[1].text = str(item.get('beneficiarios_instituciones', 0))
        
        row = table.add_row()
        row.cells[0].text = "Responsables"
        row.cells[1].text = str(item.get('responsables', '-'))
        
        row = table.add_row()
        row.cells[0].text = "Colaboradores"
        row.cells[1].text = str(item.get('colaboradores', '-'))
        
        # Espacio entre tablas
        doc.add_paragraph()
    
    # ========== SECCIÓN 2: TABLA GENERAL DE DETALLES DE ACTIVIDADES ==========
    if todas_actividades_item and todas_actividades_item.get('actividades'):
        doc.add_page_break()
        doc.add_heading('Detalles de Actividades', level=2)
        
        table = doc.add_table(rows=1, cols=7)
        format_table_word(table)
        
        # Encabezados
        headers = ['Nombre', 'Fecha', 'Estado', 'Tipo', 'Comunidad', 'Responsable', 'Beneficiarios']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Datos
        for actividad in todas_actividades_item['actividades']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(actividad.get('nombre', '-'))
            row_cells[1].text = format_date(actividad.get('fecha', '-'))
            row_cells[2].text = str(actividad.get('estado', '-'))
            row_cells[3].text = str(actividad.get('tipo_actividad', '-'))
            row_cells[4].text = str(actividad.get('comunidad', '-'))
            row_cells[5].text = str(actividad.get('responsable', '-'))
            row_cells[6].text = str(actividad.get('total_beneficiarios', 0))


def generate_word_beneficiarios_region_comunidad(doc, data):
    """Genera contenido Word para beneficiarios por región/comunidad - Tabla general con saltos por comunidad"""
    if not data or len(data) == 0:
        doc.add_paragraph('No se encontraron datos para este reporte.')
        return
    
    # Agrupar beneficiarios por comunidad
    # En este reporte, cada elemento en data es un beneficiario individual
    comunidades_map = {}
    for beneficiario in data:
        comunidad_key = beneficiario.get('comunidad', 'Sin comunidad')
        if comunidad_key not in comunidades_map:
            comunidades_map[comunidad_key] = {
                'comunidad': comunidad_key,
                'region': beneficiario.get('region', '-'),
                'beneficiarios': []
            }
        comunidades_map[comunidad_key]['beneficiarios'].append(beneficiario)
    
    if not comunidades_map:
        doc.add_paragraph('No se encontraron beneficiarios para este reporte.')
        return
    
    # Crear tabla general
    doc.add_heading('Beneficiarios por Comunidad o Región', level=2)
    
    table = doc.add_table(rows=1, cols=6)
    format_table_word(table)
    
    # Encabezados
    headers = ['Nombre', 'Tipo', 'Comunidad', 'Evento', 'DPI/Documento', 'Teléfono']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Agregar beneficiarios agrupados por comunidad
    for comunidad_key in sorted(comunidades_map.keys()):
        comunidad_data = comunidades_map[comunidad_key]
        
        # Fila de encabezado de comunidad (fila destacada)
        row = table.add_row()
        row_cells = row.cells
        
        # Primera celda con el nombre de la comunidad
        row_cells[0].text = f"{comunidad_data['comunidad']} ({comunidad_data['region']}) - {len(comunidad_data['beneficiarios'])} beneficiario{'s' if len(comunidad_data['beneficiarios']) != 1 else ''}"
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Combinar todas las celdas (de 0 a 5)
        for i in range(1, 6):
            row_cells[0].merge(row_cells[i])
        
        # Aplicar fondo gris claro al encabezado de comunidad
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D0D0D0')  # Gris un poco más oscuro
        row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        
        # Cambiar el color del texto a negro
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Agregar beneficiarios de esta comunidad
        for beneficiario in comunidad_data['beneficiarios']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(beneficiario.get('nombre', '-'))
            row_cells[1].text = str(beneficiario.get('tipo', '-'))
            row_cells[2].text = str(beneficiario.get('comunidad', '-'))
            row_cells[3].text = str(beneficiario.get('evento', '-'))
            # DPI/Documento
            row_cells[4].text = str(beneficiario.get('dpi', beneficiario.get('documento', '-')))
            # Teléfono
            row_cells[5].text = str(beneficiario.get('telefono', '-'))


def generate_word_actividad_personal(doc, data):
    """Genera contenido Word para actividad de personal - Una tabla por usuario, una hoja por tabla"""
    if not data or len(data) == 0:
        doc.add_paragraph('No se encontraron datos para este reporte.')
        return
    
    for idx, item in enumerate(data):
        # Salto de página antes de cada usuario (excepto el primero)
        if idx > 0:
            doc.add_page_break()
        
        # Agregar espacio arriba para que los datos empiecen más abajo
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Información general del usuario (fuera de la tabla, bien espaciada)
        heading = doc.add_heading(f"{item.get('nombre', 'N/A')}", level=2)
        
        # Información general con mejor formato y colores
        info_para = doc.add_paragraph()
        info_para.add_run("Puesto: ").bold = True
        info_para.add_run(f"{item.get('puesto', '-')}").font.color.rgb = RGBColor(7, 114, 210)
        
        info_para = doc.add_paragraph()
        info_para.add_run("Teléfono: ").bold = True
        info_para.add_run(f"{item.get('telefono', '-')}")
        
        info_para = doc.add_paragraph()
        info_para.add_run("Total de Eventos: ").bold = True
        info_para.add_run(f"{item.get('total_eventos', 0)}").font.color.rgb = RGBColor(7, 114, 210)
        
        info_para = doc.add_paragraph()
        info_para.add_run("Total de Avances: ").bold = True
        info_para.add_run(f"{item.get('total_avances', 0)}").font.color.rgb = RGBColor(7, 114, 210)
        
        # Espacio antes de la tabla
        doc.add_paragraph()
        
        # Crear tabla unificada con todos los eventos
        if item.get('eventos') and len(item['eventos']) > 0:
            # Título predefinido de la tabla
            title_para = doc.add_paragraph()
            title_run = title_para.add_run("Tabla de Avances Realizados en Proyectos")
            title_run.bold = True
            title_run.font.color.rgb = RGBColor(7, 114, 210)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            
            table = doc.add_table(rows=1, cols=7)
            format_table_word(table)
            
            # Encabezados
            headers = ['Evento', 'Estado', 'Tipo', 'Comunidad', 'Tipo Comunidad', 'Total Avances', 'Fecha Último Avance']
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Agregar eventos (unificados en una sola tabla)
            for evento in item['eventos']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(evento.get('nombre', '-'))
                row_cells[1].text = str(evento.get('estado', '-'))
                row_cells[2].text = str(evento.get('tipo', '-'))
                row_cells[3].text = str(evento.get('comunidad', '-'))
                row_cells[4].text = str(evento.get('tipo_comunidad', '-'))
                row_cells[5].text = str(evento.get('total_avances', 0))
                row_cells[6].text = str(evento.get('fecha_ultimo_avance', '-'))
        else:
            # Si no tiene eventos, mostrar mensaje
            doc.add_paragraph('Este colaborador no tiene eventos registrados.')


def generate_word_avances_eventos(doc, data):
    """Genera contenido Word para avances de eventos - Una hoja por evento, tabla de avances, evidencias en hoja separada"""
    if not data or len(data) == 0:
        doc.add_paragraph('No se encontraron datos para este reporte.')
        return
    
    # Agrupar avances por evento
    eventos_agrupados = {}
    for item in data:
        evento_id = item.get('evento_id', 'sin_evento')
        if evento_id not in eventos_agrupados:
            eventos_agrupados[evento_id] = {
                'evento': item.get('evento', {}),
                'comunidad': item.get('comunidad', '-'),
                'avances': []
            }
        eventos_agrupados[evento_id]['avances'].append(item)
    
    # Procesar cada evento
    for idx, (evento_id, evento_data) in enumerate(eventos_agrupados.items()):
        # Salto de página antes de cada evento (excepto el primero)
        if idx > 0:
            doc.add_page_break()
        
        # Agregar espacio arriba
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Información general del evento (fuera de la tabla, bien espaciada)
        evento_info = evento_data['evento']
        heading = doc.add_heading(f"{evento_info.get('nombre', 'N/A')}", level=2)
        
        info_para = doc.add_paragraph()
        info_para.add_run("Estado: ").bold = True
        info_para.add_run(f"{evento_info.get('estado', '-')}")
        
        info_para = doc.add_paragraph()
        info_para.add_run("Tipo: ").bold = True
        info_para.add_run(f"{evento_info.get('tipo', '-')}").font.color.rgb = RGBColor(7, 114, 210)
        
        info_para = doc.add_paragraph()
        info_para.add_run("Comunidad: ").bold = True
        info_para.add_run(f"{evento_data.get('comunidad', '-')}")
        
        info_para = doc.add_paragraph()
        info_para.add_run("Total de Avances: ").bold = True
        info_para.add_run(f"{len(evento_data['avances'])}").font.color.rgb = RGBColor(7, 114, 210)
        
        # Espacio antes de la tabla
        doc.add_paragraph()
        
        # Título predefinido de la tabla
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("Tabla de Avances Realizados")
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(7, 114, 210)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Crear tabla unificada con todos los avances
        if evento_data['avances'] and len(evento_data['avances']) > 0:
            table = doc.add_table(rows=1, cols=4)
            format_table_word(table)
            
            # Encabezados
            headers = ['Descripción', 'Responsable', 'Fecha y Hora', 'Comunidad donde se realizó el avance']
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Agregar avances
            evidencias_imagenes = []  # Para almacenar evidencias de imágenes por separado
            for avance in evento_data['avances']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(avance.get('descripcion_cambio', '-'))
                row_cells[1].text = str(avance.get('colaborador_nombre', '-'))
                row_cells[2].text = str(avance.get('fecha_display', avance.get('fecha_cambio', '-')))
                row_cells[3].text = str(avance.get('comunidad_avance', 'No se registró una comunidad para este avance'))
                
                # Recolectar evidencias de imágenes para hoja separada
                if avance.get('evidencias'):
                    for evidencia in avance['evidencias']:
                        if evidencia.get('es_imagen'):
                            evidencias_imagenes.append({
                                'titulo': avance.get('descripcion_cambio', 'Sin título'),
                                'url': evidencia.get('url', ''),
                                'nombre': evidencia.get('nombre', ''),
                                'descripcion': evidencia.get('descripcion', '')
                            })
            
            # Si hay evidencias de imágenes, crear hoja separada
            if evidencias_imagenes:
                doc.add_page_break()
                doc.add_heading('Evidencias', level=2)
                
                for evidencia in evidencias_imagenes:
                    # Título del avance
                    titulo_para = doc.add_paragraph()
                    titulo_run = titulo_para.add_run(evidencia['titulo'])
                    titulo_run.bold = True
                    titulo_run.font.color.rgb = RGBColor(7, 114, 210)
                    titulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                    
                    # Intentar agregar la imagen (5x5 cm = 1.97 inches)
                    try:
                        import requests
                        from io import BytesIO
                        
                        # Obtener la URL
                        url = evidencia['url']
                        
                        # Determinar si es una ruta local o URL
                        image_stream = None
                        if url.startswith('http'):
                            # URL externa - descargar
                            try:
                                response = requests.get(url, timeout=10)
                                image_stream = BytesIO(response.content)
                            except:
                                image_stream = None
                        else:
                            # Ruta local - construir ruta completa
                            # Remover MEDIA_URL si está presente
                            if settings.MEDIA_URL in url:
                                url = url.replace(settings.MEDIA_URL, '').lstrip('/')
                            elif url.startswith('/'):
                                url = url.lstrip('/')
                            
                            # Construir ruta completa
                            image_path = os.path.join(settings.MEDIA_ROOT, url)
                            
                            # Verificar si existe
                            if os.path.exists(image_path):
                                image_stream = open(image_path, 'rb')
                            else:
                                image_stream = None
                        
                        if image_stream:
                            para_img = doc.add_paragraph()
                            para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run_img = para_img.add_run()
                            run_img.add_picture(image_stream, width=Inches(1.97), height=Inches(1.97))
                            if hasattr(image_stream, 'close'):
                                image_stream.close()
                        else:
                            # Si no se puede cargar la imagen, mostrar el nombre del archivo
                            doc.add_paragraph(f"Imagen: {evidencia.get('nombre', 'No disponible')}")
                    except Exception as e:
                        # Si no se puede cargar la imagen, mostrar el nombre del archivo
                        doc.add_paragraph(f"Imagen: {evidencia.get('nombre', 'No disponible')}")
                    
                    doc.add_paragraph()  # Espacio entre imágenes
                    
                    # Si hay más evidencias, agregar salto de página
                    if evidencias_imagenes.index(evidencia) < len(evidencias_imagenes) - 1:
                        doc.add_page_break()
            
            # Agregar evidencias no-imagen en la tabla o en una sección separada
            evidencias_no_imagen = []
            for avance in evento_data['avances']:
                if avance.get('evidencias'):
                    for evidencia in avance['evidencias']:
                        if not evidencia.get('es_imagen'):
                            evidencias_no_imagen.append({
                                'avance': avance.get('descripcion_cambio', 'Sin título'),
                                'nombre': evidencia.get('nombre', ''),
                                'url': evidencia.get('url', '')
                            })
            
            if evidencias_no_imagen:
                if evidencias_imagenes:
                    doc.add_page_break()
                else:
                    doc.add_paragraph()
                doc.add_heading('Evidencias (Archivos)', level=2)
                
                for evidencia in evidencias_no_imagen:
                    doc.add_paragraph(f"Avance: {evidencia['avance']}")
                    doc.add_paragraph(f"Archivo: {evidencia['nombre']}")
                    doc.add_paragraph()
        else:
            doc.add_paragraph('Este evento no tiene avances registrados.')


def generate_word_evento_individual(doc, data):
    """Genera contenido Word para evento individual - Estructura completa con secciones"""
    print(f'🔵 generate_word_evento_individual llamado con data: {type(data)}')
    print(f'🔵 Keys en data: {list(data.keys()) if isinstance(data, dict) else "No es dict"}')
    
    if not data:
        doc.add_paragraph('No se encontraron datos para este reporte.')
        return
    
    evento = data.get('evento', data)  # Puede venir como data.evento o directamente como data
    print(f'🔵 Evento extraído. Keys: {list(evento.keys()) if isinstance(evento, dict) else "No es dict"}')
    print(f'🔵 Evento tiene portada: {bool(evento.get("portada"))}')
    print(f'🔵 Evento tiene cambios: {bool(evento.get("cambios"))}')
    print(f'🔵 Evento tiene evidencias: {bool(evento.get("evidencias"))}')
    
    # ========== SECCIÓN 1: TÍTULO, FOTO DE PORTADA Y DATOS GENERALES EN LA MISMA HOJA ==========
    # Agregar título del reporte
    title = doc.add_heading('Reporte de Evento Individual', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Agregar fecha de generación
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(102, 102, 102)
    doc.add_paragraph()  # Espacio
    
    # FOTO DE PORTADA (en la misma hoja)
    if evento.get('portada') and evento['portada'].get('url'):
        try:
            from io import BytesIO
            
            portada_url = evento['portada']['url']
            print(f'📷 Intentando cargar portada desde: {portada_url}')
            
            # Normalizar la URL para obtener la ruta del archivo
            relative_path = portada_url.strip()
            
            # Remover prefijos comunes de MEDIA_URL
            media_url = getattr(settings, 'MEDIA_URL', '') or ''
            posibles_prefijos = [media_url, '/media/', 'media/']
            for prefijo in posibles_prefijos:
                if prefijo and relative_path.startswith(prefijo):
                    relative_path = relative_path[len(prefijo):]
                    break
            
            # Limpiar barras iniciales
            relative_path = relative_path.lstrip('/')
            
            # Construir ruta completa del archivo
            image_path = os.path.join(settings.MEDIA_ROOT, relative_path)
            print(f'📷 Ruta del archivo: {image_path}')
            print(f'📷 ¿Existe el archivo? {os.path.exists(image_path)}')
            
            image_stream = None
            if os.path.exists(image_path):
                try:
                    image_stream = open(image_path, 'rb')
                    print(f'✅ Archivo de portada abierto correctamente')
                except Exception as e:
                    print(f'❌ Error al abrir archivo: {e}')
                    image_stream = None
            else:
                print(f'⚠️ Archivo de portada no encontrado en: {image_path}')
            
            if image_stream:
                para_img = doc.add_paragraph()
                para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = para_img.add_run()
                # Imagen rectangular ancha: width=6 inches, height=3 inches (proporción 2:1)
                run_img.add_picture(image_stream, width=Inches(6), height=Inches(3))
                if hasattr(image_stream, 'close'):
                    image_stream.close()
                print(f'✅ Imagen de portada agregada al documento')
                doc.add_paragraph()  # Espacio después de la imagen
            else:
                print(f'⚠️ No se pudo cargar la imagen de portada')
        except Exception as e:
            import traceback
            print(f'❌ Error al cargar portada: {e}')
            traceback.print_exc()
    
    # DATOS GENERALES Y DESCRIPCIÓN (en la misma hoja)
    doc.add_heading('Información General del Proyecto', level=2)
    
    # Información general con mejor formato y colores
    info_para = doc.add_paragraph()
    info_para.add_run("Nombre: ").bold = True
    info_para.add_run(f"{evento.get('nombre', '-')}").font.color.rgb = RGBColor(7, 114, 210)
    
    info_para = doc.add_paragraph()
    info_para.add_run("Tipo: ").bold = True
    info_para.add_run(f"{evento.get('tipo', '-')}").font.color.rgb = RGBColor(7, 114, 210)
    
    info_para = doc.add_paragraph()
    info_para.add_run("Estado: ").bold = True
    info_para.add_run(f"{evento.get('estado', '-')}")
    
    info_para = doc.add_paragraph()
    info_para.add_run("Fecha: ").bold = True
    info_para.add_run(f"{format_date(evento.get('fecha', evento.get('fecha_display', '-')))}")
    
    info_para = doc.add_paragraph()
    info_para.add_run("Ubicación: ").bold = True
    info_para.add_run(f"{evento.get('ubicacion', evento.get('comunidad', '-'))}")
    
    info_para = doc.add_paragraph()
    info_para.add_run("Responsable: ").bold = True
    info_para.add_run(f"{evento.get('responsable', '-')}")
    
    # Descripción
    if evento.get('descripcion'):
        doc.add_paragraph()
        desc_para = doc.add_paragraph()
        desc_para.add_run("Descripción: ").bold = True
        desc_para.add_run(f"{evento.get('descripcion', '-')}")
    
    # ========== SECCIÓN 3: PERSONAL ASIGNADO Y BENEFICIARIOS ==========
    doc.add_page_break()
    
    # Personal asignado
    if evento.get('personal') and len(evento['personal']) > 0:
        doc.add_heading('Personal Asignado', level=2)
        for persona in evento['personal']:
            personal_para = doc.add_paragraph()
            personal_para.add_run(f"• {persona.get('nombre', persona.get('username', '-'))}")
            if persona.get('puesto'):
                personal_para.add_run(f" - {persona.get('puesto', '')}")
        doc.add_paragraph()
    else:
        doc.add_heading('Personal Asignado', level=2)
        doc.add_paragraph('No hay personal asignado a este evento.')
        doc.add_paragraph()
    
    # Tabla de beneficiarios (en la misma hoja)
    if evento.get('beneficiarios') and len(evento['beneficiarios']) > 0:
        doc.add_heading('Beneficiarios', level=2)
        table = doc.add_table(rows=1, cols=4)
        format_table_word(table)
        
        headers = ['Nombre', 'Tipo', 'Comunidad', 'Información Adicional']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        for beneficiario in evento['beneficiarios']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(beneficiario.get('nombre', '-'))
            row_cells[1].text = str(beneficiario.get('tipo_display', beneficiario.get('tipo', '-')))
            # Obtener comunidad de detalles
            comunidad_benef = '-'
            if beneficiario.get('detalles') and isinstance(beneficiario['detalles'], dict):
                comunidad_benef = beneficiario['detalles'].get('comunidad_nombre', '-')
            row_cells[2].text = comunidad_benef
            row_cells[3].text = str(beneficiario.get('info_adicional', '-'))
    else:
        doc.add_heading('Beneficiarios', level=2)
        doc.add_paragraph('No hay beneficiarios registrados para este evento.')
    
    # ========== SECCIÓN 4: CAMBIOS/AVANCES (SIN EVIDENCIAS) ==========
    if evento.get('cambios') and len(evento['cambios']) > 0:
        doc.add_page_break()
        doc.add_heading('Cambios y Avances Realizados', level=2)
        
        table = doc.add_table(rows=1, cols=4)
        format_table_word(table)
        
        headers = ['Descripción', 'Responsable', 'Comunidad donde se realizó el avance', 'Fecha y Hora']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        for cambio in evento['cambios']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(cambio.get('descripcion', '-'))
            row_cells[1].text = str(cambio.get('responsable', cambio.get('responsables_display', '-')))
            # Obtener comunidades del cambio
            comunidades_cambio = cambio.get('comunidades', '')
            if not comunidades_cambio or comunidades_cambio == '':
                comunidades_cambio = 'No hay dato disponible'
            row_cells[2].text = str(comunidades_cambio)
            row_cells[3].text = str(cambio.get('fecha_display', cambio.get('fecha_cambio', '-')))
    
    # ========== SECCIÓN 5: EVIDENCIAS (SOLO FOTOS) ==========
    # Recolectar todas las evidencias de imágenes de los cambios
    evidencias_imagenes = []
    if evento.get('cambios'):
        for cambio in evento['cambios']:
            if cambio.get('evidencias'):
                for evidencia in cambio['evidencias']:
                    # Solo agregar si es imagen
                    if evidencia.get('es_imagen') or (evidencia.get('tipo') and evidencia.get('tipo', '').startswith('image/')):
                        evidencias_imagenes.append({
                            'url': evidencia.get('url', ''),
                            'nombre': evidencia.get('nombre', ''),
                            'descripcion': evidencia.get('descripcion', ''),
                            'cambio_descripcion': cambio.get('descripcion', 'Sin descripción')
                        })
    
    if evidencias_imagenes:
        doc.add_page_break()
        doc.add_heading('Evidencias (Fotografías)', level=2)
        doc.add_paragraph()  # Espacio después del título
        
        # Agrupar evidencias: hasta 2 por hoja
        for i in range(0, len(evidencias_imagenes), 2):
            # Si no es la primera iteración, agregar salto de página
            if i > 0:
                doc.add_page_break()
            
            # Primera evidencia
            if i < len(evidencias_imagenes):
                evidencia1 = evidencias_imagenes[i]
                
                # Título del avance al que pertenece
                titulo_para1 = doc.add_paragraph()
                titulo_run1 = titulo_para1.add_run(evidencia1['cambio_descripcion'])
                titulo_run1.bold = True
                titulo_run1.font.color.rgb = RGBColor(7, 114, 210)
                titulo_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Espacio después del título
                
                # Intentar agregar la imagen (5x5 cm = 1.97 inches)
                try:
                    url1 = evidencia1['url']
                    # Normalizar la URL para obtener la ruta del archivo
                    relative_path1 = url1.strip()
                    
                    # Remover prefijos comunes de MEDIA_URL
                    media_url = getattr(settings, 'MEDIA_URL', '') or ''
                    posibles_prefijos = [media_url, '/media/', 'media/']
                    for prefijo in posibles_prefijos:
                        if prefijo and relative_path1.startswith(prefijo):
                            relative_path1 = relative_path1[len(prefijo):]
                            break
                    
                    # Limpiar barras iniciales
                    relative_path1 = relative_path1.lstrip('/')
                    
                    # Construir ruta completa del archivo
                    image_path1 = os.path.join(settings.MEDIA_ROOT, relative_path1)
                    print(f'📷 Evidencia 1: Intentando cargar desde: {image_path1}')
                    
                    image_stream1 = None
                    if os.path.exists(image_path1):
                        try:
                            image_stream1 = open(image_path1, 'rb')
                            print(f'✅ Evidencia 1: Archivo abierto correctamente')
                        except Exception as e:
                            print(f'❌ Evidencia 1: Error al abrir archivo: {e}')
                            image_stream1 = None
                    else:
                        print(f'⚠️ Evidencia 1: Archivo no encontrado en: {image_path1}')
                    
                    if image_stream1:
                        para_img1 = doc.add_paragraph()
                        para_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_img1 = para_img1.add_run()
                        run_img1.add_picture(image_stream1, width=Inches(1.97), height=Inches(1.97))
                        if hasattr(image_stream1, 'close'):
                            image_stream1.close()
                        print(f'✅ Evidencia 1: Imagen agregada al documento')
                    else:
                        doc.add_paragraph(f"Imagen: {evidencia1.get('nombre', 'No disponible')}")
                except Exception as e:
                    import traceback
                    print(f'❌ Evidencia 1: Error al cargar imagen: {e}')
                    traceback.print_exc()
                    doc.add_paragraph(f"Imagen: {evidencia1.get('nombre', 'No disponible')}")
                
                doc.add_paragraph()  # Espacio después de la imagen
            
            # Segunda evidencia (si existe)
            if i + 1 < len(evidencias_imagenes):
                evidencia2 = evidencias_imagenes[i + 1]
                
                # Título del avance al que pertenece
                titulo_para2 = doc.add_paragraph()
                titulo_run2 = titulo_para2.add_run(evidencia2['cambio_descripcion'])
                titulo_run2.bold = True
                titulo_run2.font.color.rgb = RGBColor(7, 114, 210)
                titulo_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Espacio después del título
                
                # Intentar agregar la imagen (5x5 cm = 1.97 inches)
                try:
                    url2 = evidencia2['url']
                    # Normalizar la URL para obtener la ruta del archivo
                    relative_path2 = url2.strip()
                    
                    # Remover prefijos comunes de MEDIA_URL
                    media_url = getattr(settings, 'MEDIA_URL', '') or ''
                    posibles_prefijos = [media_url, '/media/', 'media/']
                    for prefijo in posibles_prefijos:
                        if prefijo and relative_path2.startswith(prefijo):
                            relative_path2 = relative_path2[len(prefijo):]
                            break
                    
                    # Limpiar barras iniciales
                    relative_path2 = relative_path2.lstrip('/')
                    
                    # Construir ruta completa del archivo
                    image_path2 = os.path.join(settings.MEDIA_ROOT, relative_path2)
                    print(f'📷 Evidencia 2: Intentando cargar desde: {image_path2}')
                    
                    image_stream2 = None
                    if os.path.exists(image_path2):
                        try:
                            image_stream2 = open(image_path2, 'rb')
                            print(f'✅ Evidencia 2: Archivo abierto correctamente')
                        except Exception as e:
                            print(f'❌ Evidencia 2: Error al abrir archivo: {e}')
                            image_stream2 = None
                    else:
                        print(f'⚠️ Evidencia 2: Archivo no encontrado en: {image_path2}')
                    
                    if image_stream2:
                        para_img2 = doc.add_paragraph()
                        para_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_img2 = para_img2.add_run()
                        run_img2.add_picture(image_stream2, width=Inches(1.97), height=Inches(1.97))
                        if hasattr(image_stream2, 'close'):
                            image_stream2.close()
                        print(f'✅ Evidencia 2: Imagen agregada al documento')
                    else:
                        doc.add_paragraph(f"Imagen: {evidencia2.get('nombre', 'No disponible')}")
                except Exception as e:
                    import traceback
                    print(f'❌ Evidencia 2: Error al cargar imagen: {e}')
                    traceback.print_exc()
                    doc.add_paragraph(f"Imagen: {evidencia2.get('nombre', 'No disponible')}")
                
                doc.add_paragraph()  # Espacio después de la imagen
    
    # ========== SECCIÓN 6: GALERÍA DE IMÁGENES (ANEXOS) ==========
    # Obtener imágenes de la galería del evento (eventos_galeria)
    galeria_imagenes = []
    if evento.get('evidencias'):  # evidencias contiene las imágenes de la galería
        for imagen in evento['evidencias']:
            # Solo agregar si es imagen
            if imagen.get('es_imagen') or (imagen.get('tipo') and imagen.get('tipo', '').startswith('image/')):
                galeria_imagenes.append({
                    'url': imagen.get('url', ''),
                    'nombre': imagen.get('nombre', ''),
                    'descripcion': imagen.get('descripcion', '')
                })
    
    if galeria_imagenes:
        doc.add_page_break()
        doc.add_heading('Anexos - Galería de Imágenes', level=2)
        doc.add_paragraph()  # Espacio
        
        # Mostrar 2 imágenes por línea
        for i in range(0, len(galeria_imagenes), 2):
            # Crear una tabla de 2 columnas para alinear las imágenes
            table_galeria = doc.add_table(rows=1, cols=2)
            table_galeria.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Primera imagen
            if i < len(galeria_imagenes):
                imagen1 = galeria_imagenes[i]
                cell1 = table_galeria.rows[0].cells[0]
                para1 = cell1.paragraphs[0]
                para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                try:
                    url1 = imagen1['url']
                    # Normalizar la URL
                    relative_path1 = url1.strip()
                    media_url = getattr(settings, 'MEDIA_URL', '') or ''
                    posibles_prefijos = [media_url, '/media/', 'media/']
                    for prefijo in posibles_prefijos:
                        if prefijo and relative_path1.startswith(prefijo):
                            relative_path1 = relative_path1[len(prefijo):]
                            break
                    relative_path1 = relative_path1.lstrip('/')
                    image_path1 = os.path.join(settings.MEDIA_ROOT, relative_path1)
                    
                    if os.path.exists(image_path1):
                        image_stream1 = open(image_path1, 'rb')
                        run1 = para1.add_run()
                        run1.add_picture(image_stream1, width=Inches(1.97), height=Inches(1.97))  # 5x5 cm
                        image_stream1.close()
                except Exception as e:
                    para1.add_run(f"Imagen: {imagen1.get('nombre', 'No disponible')}")
                    print(f'⚠️ Error al cargar imagen de galería: {e}')
            
            # Segunda imagen (si existe)
            if i + 1 < len(galeria_imagenes):
                imagen2 = galeria_imagenes[i + 1]
                cell2 = table_galeria.rows[0].cells[1]
                para2 = cell2.paragraphs[0]
                para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                try:
                    url2 = imagen2['url']
                    # Normalizar la URL
                    relative_path2 = url2.strip()
                    media_url = getattr(settings, 'MEDIA_URL', '') or ''
                    posibles_prefijos = [media_url, '/media/', 'media/']
                    for prefijo in posibles_prefijos:
                        if prefijo and relative_path2.startswith(prefijo):
                            relative_path2 = relative_path2[len(prefijo):]
                            break
                    relative_path2 = relative_path2.lstrip('/')
                    image_path2 = os.path.join(settings.MEDIA_ROOT, relative_path2)
                    
                    if os.path.exists(image_path2):
                        image_stream2 = open(image_path2, 'rb')
                        run2 = para2.add_run()
                        run2.add_picture(image_stream2, width=Inches(1.97), height=Inches(1.97))  # 5x5 cm
                        image_stream2.close()
                except Exception as e:
                    para2.add_run(f"Imagen: {imagen2.get('nombre', 'No disponible')}")
                    print(f'⚠️ Error al cargar imagen de galería: {e}')
            
            doc.add_paragraph()  # Espacio entre filas
    

def generate_word_comunidades(doc, data):
    """Genera contenido Word para reporte de comunidades"""
    # Manejar caso donde data viene como dict con 'comunidades'
    if isinstance(data, dict) and 'comunidades' in data:
        data = data['comunidades']
    
    if not data or (isinstance(data, list) and len(data) == 0):
        doc.add_paragraph('No se encontraron datos para este reporte.')
        return
    
    # Asegurar que data es una lista
    if not isinstance(data, list):
        doc.add_paragraph('Error: formato de datos inválido para el reporte de comunidades.')
        return
    
    for idx, item in enumerate(data):
        # Validar que item es un diccionario
        if not isinstance(item, dict):
            continue
        
        # Salto de página antes de cada comunidad (excepto la primera)
        if idx > 0:
            doc.add_page_break()
        
        # ========== SECCIÓN 1: INFORMACIÓN GENERAL DE LA COMUNIDAD ==========
        # Título de la comunidad
        heading = doc.add_heading(f"{item.get('nombre', 'N/A')}", level=2)
        
        # Espacio después del título
        doc.add_paragraph()
        
        # Información general con formato mejorado (todos en la misma hoja)
        # COCODE
        info_para = doc.add_paragraph()
        info_para.add_run("COCODE: ").bold = True
        info_para.add_run(f"{item.get('cocode', '-')}").font.color.rgb = RGBColor(7, 114, 210)
        
        # Región
        info_para = doc.add_paragraph()
        info_para.add_run("Región: ").bold = True
        info_para.add_run(f"{item.get('region', '-')}").font.color.rgb = RGBColor(7, 114, 210)
        
        # Tipo
        info_para = doc.add_paragraph()
        info_para.add_run("Tipo: ").bold = True
        info_para.add_run(f"{item.get('tipo', '-')}").font.color.rgb = RGBColor(7, 114, 210)
        
        # Número de Beneficiarios
        info_para = doc.add_paragraph()
        info_para.add_run("Número de Beneficiarios: ").bold = True
        total_beneficiarios = item.get('numero_beneficiarios', item.get('total_beneficiarios', 0))
        info_para.add_run(f"{total_beneficiarios}").font.color.rgb = RGBColor(7, 114, 210)
        
        # Número de Proyectos
        info_para = doc.add_paragraph()
        info_para.add_run("Número de Proyectos: ").bold = True
        total_proyectos = item.get('numero_proyectos', item.get('total_proyectos', 0))
        info_para.add_run(f"{total_proyectos}").font.color.rgb = RGBColor(7, 114, 210)
        
        doc.add_paragraph()  # Espacio adicional
        
        # ========== SECCIÓN 2: TABLA DE PROYECTOS/EVENTOS ==========
        if item.get('proyectos') and len(item['proyectos']) > 0:
            # Título de la sección de proyectos
            doc.add_heading('Proyectos/Eventos', level=3)
            doc.add_paragraph()  # Espacio antes de la tabla
            
            # Crear tabla de proyectos
            table = doc.add_table(rows=1, cols=4)
            format_table_word(table)
            
            # Encabezados
            headers = ['Nombre', 'Tipo', 'Estado', 'Fecha']
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Agregar filas de proyectos
            for proyecto in item['proyectos']:
                if isinstance(proyecto, dict):
                    row_cells = table.add_row().cells
                    row_cells[0].text = proyecto.get('nombre', '-')
                    row_cells[1].text = proyecto.get('tipo', '-')
                    row_cells[2].text = proyecto.get('estado', '-')
                    row_cells[3].text = format_date(proyecto.get('fecha', '-'))
            
            doc.add_paragraph()  # Espacio después de la tabla


def generate_word_actividad_usuarios(doc, data):
    """Genera contenido Word para actividad de usuarios"""
    # Manejar caso donde data viene como dict con 'usuarios'
    if isinstance(data, dict) and 'usuarios' in data:
        data = data['usuarios']
    
    if not data or (isinstance(data, list) and len(data) == 0):
        doc.add_paragraph('No se encontraron datos para este reporte.')
        return
    
    # Asegurar que data es una lista
    if not isinstance(data, list):
        doc.add_paragraph('Error: formato de datos inválido para el reporte de actividad de usuarios.')
        return
    
    for idx, item in enumerate(data):
        # Validar que item es un diccionario
        if not isinstance(item, dict):
            continue
        
        # Salto de página antes de cada usuario (excepto el primero)
        if idx > 0:
            doc.add_page_break()
        
        # ========== SECCIÓN: INFORMACIÓN DEL USUARIO ==========
        # Título del usuario
        nombre_usuario = item.get('nombre', item.get('username', 'N/A'))
        heading = doc.add_heading(f"Usuario: {nombre_usuario}", level=2)
        
        # Espacio después del título
        doc.add_paragraph()
        
        # Información general del usuario con formato mejorado
        info_para = doc.add_paragraph()
        info_para.add_run("Usuario: ").bold = True
        info_para.add_run(f"{item.get('username', '-')}").font.color.rgb = RGBColor(7, 114, 210)
        
        info_para = doc.add_paragraph()
        info_para.add_run("Rol: ").bold = True
        info_para.add_run(f"{item.get('rol_display', item.get('rol', '-'))}").font.color.rgb = RGBColor(7, 114, 210)
        
        if item.get('puesto_nombre') or item.get('puesto'):
            info_para = doc.add_paragraph()
            info_para.add_run("Puesto: ").bold = True
            puesto = item.get('puesto_nombre', item.get('puesto', '-'))
            info_para.add_run(f"{puesto}").font.color.rgb = RGBColor(7, 114, 210)
        
        info_para = doc.add_paragraph()
        info_para.add_run("Total de cambios: ").bold = True
        info_para.add_run(f"{item.get('total_cambios', 0)}").font.color.rgb = RGBColor(7, 114, 210)
        
        doc.add_paragraph()  # Espacio adicional
        
        # ========== SECCIÓN: TABLA DE CAMBIOS POR PROYECTO ==========
        if item.get('cambios') and len(item['cambios']) > 0:
            # Agrupar cambios por evento/proyecto
            cambios_por_evento = {}
            for cambio in item['cambios']:
                if not isinstance(cambio, dict):
                    continue
                
                evento_id = cambio.get('evento_id', 'sin_evento')
                evento_nombre = cambio.get('evento_nombre', 'Sin evento')
                
                if evento_id not in cambios_por_evento:
                    cambios_por_evento[evento_id] = {
                        'nombre': evento_nombre,
                        'cambios': []
                    }
                
                cambios_por_evento[evento_id]['cambios'].append(cambio)
            
            # Crear tabla general para todos los cambios
            doc.add_heading('Tabla de Cambios Realizados', level=3)
            doc.add_paragraph()  # Espacio antes de la tabla
            
            # Crear tabla con columnas: Proyecto, Tipo de Cambio, Descripción, Fecha
            table = doc.add_table(rows=1, cols=4)
            format_table_word(table)
            
            # Encabezados
            headers = ['Proyecto', 'Tipo de Cambio', 'Descripción', 'Fecha y Hora']
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Agregar filas agrupadas por proyecto
            for evento_id, evento_data in cambios_por_evento.items():
                cambios_evento = evento_data['cambios']
                
                # Agregar filas de cambios para este proyecto
                for cambio_idx, cambio in enumerate(cambios_evento):
                    row_cells = table.add_row().cells
                    
                    # En la primera fila de cada proyecto, mostrar el nombre del proyecto
                    if cambio_idx == 0:
                        # Agregar nombre del proyecto en negrita y azul
                        proyecto_para = row_cells[0].paragraphs[0]
                        # Limpiar el párrafo eliminando todos los runs existentes
                        proyecto_para.clear()
                        proyecto_run = proyecto_para.add_run(evento_data['nombre'])
                        proyecto_run.bold = True
                        proyecto_run.font.color.rgb = RGBColor(7, 114, 210)
                    else:
                        # En las siguientes filas, dejar vacío
                        row_cells[0].text = ''
                    
                    row_cells[1].text = cambio.get('tipo_cambio', '-')
                    row_cells[2].text = cambio.get('descripcion', '-')
                    # Usar fecha_display si está disponible, sino format_date de fecha
                    fecha_str = cambio.get('fecha_display', '')
                    if not fecha_str and cambio.get('fecha'):
                        fecha_str = format_date(cambio.get('fecha', '-'))
                    row_cells[3].text = fecha_str if fecha_str else '-'
            
            doc.add_paragraph()  # Espacio después de la tabla


def generate_word_reporte_general(doc, data):
    """Genera contenido Word para reporte general - Datos primero, luego tablas"""
    if not data:
        doc.add_paragraph('No se encontraron datos para este reporte.')
        return
    
    # ========== SECCIÓN 1: TODOS LOS DATOS EN UNA MISMA HOJA ==========
    doc.add_heading('Resumen General', level=2)
    
    # Tipos de eventos (dato)
    if 'tipos_eventos' in data:
        tipos = data['tipos_eventos']
        info_para = doc.add_paragraph()
        info_para.add_run("Total de Tipos de Evento: ").bold = True
        info_para.add_run(f"Capacitación: {tipos.get('capacitacion', 0)}, Entrega: {tipos.get('entrega', 0)}, Proyecto de Ayuda: {tipos.get('proyecto_ayuda', 0)}").font.color.rgb = RGBColor(7, 114, 210)
        doc.add_paragraph()
    
    # Total de beneficiarios alcanzado (dato)
    if 'total_beneficiarios' in data:
        info_para = doc.add_paragraph()
        info_para.add_run("Total de Beneficiarios Alcanzados: ").bold = True
        info_para.add_run(f"{data['total_beneficiarios']}").font.color.rgb = RGBColor(7, 114, 210)
        doc.add_paragraph()
    
    # Tipos de beneficiarios (dato)
    if 'tipos_beneficiarios' in data:
        tipos = data['tipos_beneficiarios']
        info_para = doc.add_paragraph()
        info_para.add_run("Total de Tipos de Beneficiarios: ").bold = True
        info_para.add_run(f"Individuales: {tipos.get('individual', 0)}, Familias: {tipos.get('familia', 0)}, Instituciones: {tipos.get('institucion', 0)}, Mujeres: {tipos.get('mujeres', 0)}, Hombres: {tipos.get('hombres', 0)}").font.color.rgb = RGBColor(7, 114, 210)
        doc.add_paragraph()
    
    # Total de comunidades alcanzadas (dato - solo el número)
    if 'total_comunidades' in data:
        total_comunidades = data['total_comunidades']
        info_para = doc.add_paragraph()
        info_para.add_run("Total de Comunidades Alcanzadas: ").bold = True
        info_para.add_run(f"{total_comunidades.get('total', 0)}").font.color.rgb = RGBColor(7, 114, 210)
        doc.add_paragraph()
    
    # Total de avances en Proyectos (dato)
    if 'total_avances' in data:
        info_para = doc.add_paragraph()
        info_para.add_run("Total de Avances en Proyectos: ").bold = True
        info_para.add_run(f"{data['total_avances']}").font.color.rgb = RGBColor(7, 114, 210)
        doc.add_paragraph()
    
    # Evento con más avances y cambios (pequeña tabla)
    if 'evento_mas_avances' in data and data['evento_mas_avances']:
        evento = data['evento_mas_avances']
        doc.add_heading('Evento con Más Avances y Cambios', level=3)
        table = doc.add_table(rows=1, cols=3)
        format_table_word(table)
        
        headers = ['Nombre', 'Total de Avances', 'Total de Cambios']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(evento.get('nombre', '-'))
        row_cells[1].text = str(evento.get('total_avances', 0))
        row_cells[2].text = str(evento.get('total_cambios', 0))
        doc.add_paragraph()
    
    # ========== SECCIÓN 2: TABLAS - CADA UNA EN SU PROPIA HOJA ==========
    
    # Tabla 1: Total de eventos
    if 'total_eventos' in data:
        total_eventos = data['total_eventos']
        if total_eventos.get('lista_eventos') and len(total_eventos['lista_eventos']) > 0:
            doc.add_page_break()
            doc.add_heading('Total de Eventos', level=2)
            info_para = doc.add_paragraph()
            info_para.add_run(f"Total: {total_eventos.get('total', 0)}").bold = True
            doc.add_paragraph()
            
            table = doc.add_table(rows=1, cols=6)
            format_table_word(table)
            
            headers = ['Nombre', 'Tipo', 'Estado', 'Fecha', 'Comunidad', 'Beneficiarios']
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            for evento in total_eventos['lista_eventos']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(evento.get('nombre', '-'))
                row_cells[1].text = str(evento.get('tipo', '-'))
                row_cells[2].text = str(evento.get('estado', '-'))
                row_cells[3].text = format_date(evento.get('fecha', '-'))
                row_cells[4].text = str(evento.get('comunidad', '-'))
                row_cells[5].text = str(evento.get('beneficiarios', 0))
    
    # Tabla 2: Total de comunidades alcanzadas (si tiene lista)
    if 'total_comunidades' in data:
        total_comunidades = data['total_comunidades']
        if total_comunidades.get('lista_comunidades') and len(total_comunidades['lista_comunidades']) > 0:
            doc.add_page_break()
            doc.add_heading('Total de Comunidades Alcanzadas', level=2)
            info_para = doc.add_paragraph()
            info_para.add_run(f"Total: {total_comunidades.get('total', 0)}").bold = True
            doc.add_paragraph()
            
            table = doc.add_table(rows=1, cols=2)
            format_table_word(table)
            
            headers = ['Comunidad', 'Región']
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            for comunidad in total_comunidades['lista_comunidades']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(comunidad.get('nombre', '-'))
                row_cells[1].text = str(comunidad.get('region', '-'))
    
    # Tabla 3: Comunidades con más beneficiarios y eventos
    if 'comunidades_mas_beneficiarios' in data and len(data['comunidades_mas_beneficiarios']) > 0:
        doc.add_page_break()
        doc.add_heading('Comunidades con Más Beneficiarios y Eventos', level=2)
        comunidades = data['comunidades_mas_beneficiarios']
        
        table = doc.add_table(rows=1, cols=3)
        format_table_word(table)
        
        headers = ['Comunidad', 'Total de Beneficiarios', 'Total de Eventos']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        for com in comunidades:
            row_cells = table.add_row().cells
            row_cells[0].text = str(com.get('nombre', '-'))
            row_cells[1].text = str(com.get('total_beneficiarios', 0))
            row_cells[2].text = str(com.get('total_eventos', 0))
    
    # Tabla 4: Eventos con más beneficiarios
    if 'eventos_mas_beneficiarios' in data and len(data['eventos_mas_beneficiarios']) > 0:
        doc.add_page_break()
        doc.add_heading('Eventos con Más Beneficiarios', level=2)
        eventos = data['eventos_mas_beneficiarios']
        
        table = doc.add_table(rows=1, cols=2)
        format_table_word(table)
        
        headers = ['Evento', 'Total de Beneficiarios']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        for evento in eventos:
            row_cells = table.add_row().cells
            row_cells[0].text = str(evento.get('nombre', '-'))
            row_cells[1].text = str(evento.get('total_beneficiarios', 0))


# ========== FUNCIONES ESPECÍFICAS PARA HTML/PDF ==========

def generate_html_actividades_region_comunidad(data):
    """Genera HTML para actividades por región/comunidad"""
    html_parts = []
    
    if not data or len(data) == 0:
        return '<p>No se encontraron datos para este reporte.</p>'
    
    # Separar items con resumen (sin actividades) de la entrada con todas las actividades
    items_resumen = []
    todas_actividades_item = None
    
    for item in data:
        # Si tiene actividades y es la entrada de "Todas las Actividades" o similar
        if item.get('actividades') and len(item.get('actividades', [])) > 0:
            todas_actividades_item = item
        else:
            items_resumen.append(item)
    
    # ========== SECCIÓN 1: TABLAS DE RESUMEN (2 por página) ==========
    for idx, item in enumerate(items_resumen):
        # Salto de página cada 2 items
        if idx > 0 and idx % 2 == 0:
            html_parts.append('<div class="section-break"></div>')
        
        # Tabla de resumen (2 por página, más espacio)
        html_parts.append(f'<h3 style="text-align: center; margin-bottom: 0.3cm; font-size: 11pt;">{item.get("nombre", "N/A")}</h3>')
        html_parts.append('<table style="width: 48%; display: inline-table; margin: 0.3cm 1%; vertical-align: top; font-size: 9pt; border: 1px solid #000;">')
        html_parts.append('<thead>')
        html_parts.append('<tr><th style="width: 50%;">Información</th><th style="width: 50%;">Valor</th></tr>')
        html_parts.append('</thead>')
        html_parts.append('<tbody>')
        
        html_parts.append('<tr><td style="font-weight: bold; padding: 5px;">Región:</td><td style="padding: 5px;">' + str(item.get('region', '-')) + '</td></tr>')
        html_parts.append('<tr><td style="font-weight: bold; padding: 5px;">Total actividades:</td><td style="padding: 5px; color: #0772d2;">' + str(item.get('total_actividades', 0)) + '</td></tr>')
        html_parts.append('<tr><td style="font-weight: bold; padding: 5px;">Total beneficiarios:</td><td style="padding: 5px; color: #0772d2;">' + str(item.get('total_beneficiarios', 0)) + '</td></tr>')
        html_parts.append('<tr><td style="font-weight: bold; padding: 5px;">Individuales:</td><td style="padding: 5px;">' + str(item.get('beneficiarios_individuales', 0)) + '</td></tr>')
        html_parts.append('<tr><td style="font-weight: bold; padding: 5px;">Familias:</td><td style="padding: 5px;">' + str(item.get('beneficiarios_familias', 0)) + '</td></tr>')
        html_parts.append('<tr><td style="font-weight: bold; padding: 5px;">Instituciones:</td><td style="padding: 5px;">' + str(item.get('beneficiarios_instituciones', 0)) + '</td></tr>')
        html_parts.append('<tr><td style="font-weight: bold; padding: 5px;">Responsables:</td><td style="padding: 5px; font-size: 8pt;">' + str(item.get('responsables', '-')) + '</td></tr>')
        html_parts.append('<tr><td style="font-weight: bold; padding: 5px;">Colaboradores:</td><td style="padding: 5px; font-size: 8pt;">' + str(item.get('colaboradores', '-')) + '</td></tr>')
        
        html_parts.append('</tbody></table>')
        
        # Agregar salto de línea después de cada 2 tablas
        if (idx + 1) % 2 == 0:
            html_parts.append('<div style="clear: both; page-break-after: always;"></div>')
    
    # ========== SECCIÓN 2: TABLA GENERAL DE DETALLES DE ACTIVIDADES ==========
    if todas_actividades_item and todas_actividades_item.get('actividades'):
        html_parts.append('<div class="section-break"><h2>Detalles de Actividades</h2>')
        html_parts.append('<table>')
        html_parts.append('<thead><tr><th>Nombre</th><th>Fecha</th><th>Estado</th><th>Tipo</th><th>Comunidad</th><th>Responsable</th><th>Beneficiarios</th></tr></thead>')
        html_parts.append('<tbody>')
        
        for actividad in todas_actividades_item['actividades']:
            html_parts.append('<tr>')
            html_parts.append(f'<td>{actividad.get("nombre", "-")}</td>')
            html_parts.append(f'<td>{format_date(actividad.get("fecha", "-"))}</td>')
            html_parts.append(f'<td>{actividad.get("estado", "-")}</td>')
            html_parts.append(f'<td>{actividad.get("tipo_actividad", "-")}</td>')
            html_parts.append(f'<td>{actividad.get("comunidad", "-")}</td>')
            html_parts.append(f'<td>{actividad.get("responsable", "-")}</td>')
            html_parts.append(f'<td>{actividad.get("total_beneficiarios", 0)}</td>')
            html_parts.append('</tr>')
        
        html_parts.append('</tbody></table></div>')
    
    return '\n'.join(html_parts)


def generate_html_beneficiarios_region_comunidad(data):
    """Genera HTML para beneficiarios por región/comunidad - Tabla general con saltos por comunidad"""
    html_parts = []
    
    if not data or len(data) == 0:
        return '<p>No se encontraron datos para este reporte.</p>'
    
    # Agrupar beneficiarios por comunidad
    # En este reporte, cada elemento en data es un beneficiario individual
    comunidades_map = {}
    for beneficiario in data:
        comunidad_key = beneficiario.get('comunidad', 'Sin comunidad')
        if comunidad_key not in comunidades_map:
            comunidades_map[comunidad_key] = {
                'comunidad': comunidad_key,
                'region': beneficiario.get('region', '-'),
                'beneficiarios': []
            }
        comunidades_map[comunidad_key]['beneficiarios'].append(beneficiario)
    
    if not comunidades_map:
        return '<p>No se encontraron beneficiarios para este reporte.</p>'
    
    # Crear tabla general
    html_parts.append('<h2>Beneficiarios por Comunidad o Región</h2>')
    html_parts.append('<table>')
    html_parts.append('<thead><tr><th>Nombre</th><th>Tipo</th><th>Comunidad</th><th>Evento</th><th>DPI/Documento</th><th>Teléfono</th></tr></thead>')
    html_parts.append('<tbody>')
    
    # Agregar beneficiarios agrupados por comunidad
    for comunidad_key in sorted(comunidades_map.keys()):
        comunidad_data = comunidades_map[comunidad_key]
        
        # Fila de encabezado de comunidad (fila destacada con fondo gris claro)
        html_parts.append(f'<tr style="background-color: #D0D0D0; color: #000000;">')
        html_parts.append(f'<td colspan="6" style="font-weight: bold; padding: 8px; text-align: center;">')
        html_parts.append(f'{comunidad_data["comunidad"]} ({comunidad_data["region"]}) - {len(comunidad_data["beneficiarios"])} beneficiario{"s" if len(comunidad_data["beneficiarios"]) != 1 else ""}')
        html_parts.append('</td>')
        html_parts.append('</tr>')
        
        # Agregar beneficiarios de esta comunidad
        for beneficiario in comunidad_data['beneficiarios']:
            html_parts.append('<tr>')
            html_parts.append(f'<td>{beneficiario.get("nombre", "-")}</td>')
            html_parts.append(f'<td>{beneficiario.get("tipo", "-")}</td>')
            html_parts.append(f'<td>{beneficiario.get("comunidad", "-")}</td>')
            html_parts.append(f'<td>{beneficiario.get("evento", "-")}</td>')
            # DPI/Documento
            html_parts.append(f'<td>{beneficiario.get("dpi", beneficiario.get("documento", "-"))}</td>')
            # Teléfono
            html_parts.append(f'<td>{beneficiario.get("telefono", "-")}</td>')
            html_parts.append('</tr>')
    
    html_parts.append('</tbody></table>')
    
    return '\n'.join(html_parts)


def generate_html_actividad_personal(data):
    """Genera HTML para actividad de personal - Una tabla por usuario, una hoja por tabla"""
    html_parts = []
    
    if not data or len(data) == 0:
        return '<p>No se encontraron datos para este reporte.</p>'
    
    for idx, item in enumerate(data):
        # Salto de página antes de cada usuario (excepto el primero)
        class_attr = ' class="section-break"' if idx > 0 else ''
        
        # Agregar espacio arriba para que los datos empiecen más abajo
        html_parts.append('<div style="margin-top: 2cm;"></div>')
        
        html_parts.append(f'<h2{class_attr}>{item.get("nombre", "N/A")}</h2>')
        
        # Información general del usuario (fuera de la tabla, bien espaciada)
        html_parts.append('<div class="metric-box">')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Puesto:</span> <strong style="color: #0772d2;">{item.get("puesto", "-")}</strong></p>')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Teléfono:</span> <strong>{item.get("telefono", "-")}</strong></p>')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Total de Eventos:</span> <strong style="color: #0772d2;">{item.get("total_eventos", 0)}</strong></p>')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Total de Avances:</span> <strong style="color: #0772d2;">{item.get("total_avances", 0)}</strong></p>')
        html_parts.append('</div>')
        
        # Crear tabla unificada con todos los eventos
        if item.get('eventos') and len(item['eventos']) > 0:
            # Título predefinido de la tabla
            html_parts.append('<h3 style="text-align: center; color: #0772d2; margin-top: 0.5cm; margin-bottom: 0.3cm;">Tabla de Avances Realizados en Proyectos</h3>')
            
            html_parts.append('<table>')
            html_parts.append('<thead><tr><th>Evento</th><th>Estado</th><th>Tipo</th><th>Comunidad</th><th>Tipo Comunidad</th><th>Total Avances</th><th>Fecha Último Avance</th></tr></thead>')
            html_parts.append('<tbody>')
            
            for evento in item['eventos']:
                html_parts.append('<tr>')
                html_parts.append(f'<td>{evento.get("nombre", "-")}</td>')
                html_parts.append(f'<td>{evento.get("estado", "-")}</td>')
                html_parts.append(f'<td>{evento.get("tipo", "-")}</td>')
                html_parts.append(f'<td>{evento.get("comunidad", "-")}</td>')
                html_parts.append(f'<td>{evento.get("tipo_comunidad", "-")}</td>')
                html_parts.append(f'<td>{evento.get("total_avances", 0)}</td>')
                html_parts.append(f'<td>{evento.get("fecha_ultimo_avance", "-")}</td>')
                html_parts.append('</tr>')
            
            html_parts.append('</tbody></table>')
        else:
            html_parts.append('<p>Este colaborador no tiene eventos registrados.</p>')
    
    return '\n'.join(html_parts)


def generate_html_avances_eventos(data):
    """Genera HTML para avances de eventos - Una hoja por evento, tabla de avances, evidencias en hoja separada"""
    html_parts = []
    
    if not data or len(data) == 0:
        return '<p>No se encontraron datos para este reporte.</p>'
    
    # Agrupar avances por evento
    eventos_agrupados = {}
    for item in data:
        evento_id = item.get('evento_id', 'sin_evento')
        if evento_id not in eventos_agrupados:
            eventos_agrupados[evento_id] = {
                'evento': item.get('evento', {}),
                'comunidad': item.get('comunidad', '-'),
                'avances': []
            }
        eventos_agrupados[evento_id]['avances'].append(item)
    
    # Procesar cada evento
    for idx, (evento_id, evento_data) in enumerate(eventos_agrupados.items()):
        # Salto de página antes de cada evento (excepto el primero)
        class_attr = ' class="section-break"' if idx > 0 else ''
        
        # Agregar espacio arriba
        html_parts.append('<div style="margin-top: 2cm;"></div>')
        
        # Información general del evento (fuera de la tabla, bien espaciada)
        evento_info = evento_data['evento']
        html_parts.append(f'<h2{class_attr}>{evento_info.get("nombre", "N/A")}</h2>')
        
        html_parts.append('<div class="metric-box">')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Estado:</span> <strong>{evento_info.get("estado", "-")}</strong></p>')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Tipo:</span> <strong style="color: #0772d2;">{evento_info.get("tipo", "-")}</strong></p>')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Comunidad:</span> <strong>{evento_data.get("comunidad", "-")}</strong></p>')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Total de Avances:</span> <strong style="color: #0772d2;">{len(evento_data["avances"])}</strong></p>')
        html_parts.append('</div>')
        
        # Título predefinido de la tabla
        html_parts.append('<h3 style="text-align: center; color: #0772d2; margin-top: 0.5cm; margin-bottom: 0.3cm;">Tabla de Avances Realizados</h3>')
        
        # Crear tabla unificada con todos los avances
        if evento_data['avances'] and len(evento_data['avances']) > 0:
            html_parts.append('<table>')
            html_parts.append('<thead><tr><th>Descripción</th><th>Responsable</th><th>Fecha y Hora</th><th>Comunidad donde se realizó el avance</th></tr></thead>')
            html_parts.append('<tbody>')
            
            evidencias_imagenes = []  # Para almacenar evidencias de imágenes por separado
            
            for avance in evento_data['avances']:
                html_parts.append('<tr>')
                html_parts.append(f'<td>{avance.get("descripcion_cambio", "-")}</td>')
                html_parts.append(f'<td>{avance.get("colaborador_nombre", "-")}</td>')
                html_parts.append(f'<td>{avance.get("fecha_display", avance.get("fecha_cambio", "-"))}</td>')
                html_parts.append(f'<td>{avance.get("comunidad_avance", "No se registró una comunidad para este avance")}</td>')
                html_parts.append('</tr>')
                
                # Recolectar evidencias de imágenes para hoja separada
                if avance.get('evidencias'):
                    for evidencia in avance['evidencias']:
                        if evidencia.get('es_imagen'):
                            evidencias_imagenes.append({
                                'titulo': avance.get('descripcion_cambio', 'Sin título'),
                                'url': evidencia.get('url', ''),
                                'nombre': evidencia.get('nombre', ''),
                                'descripcion': evidencia.get('descripcion', '')
                            })
            
            html_parts.append('</tbody></table>')
            
            # Si hay evidencias de imágenes, crear hoja separada
            if evidencias_imagenes:
                html_parts.append('<div class="section-break"><h2>Evidencias</h2>')
                
                for evidencia in evidencias_imagenes:
                    # Título del avance
                    html_parts.append(f'<h3 style="text-align: center; color: #0772d2; margin-top: 0.5cm; margin-bottom: 0.3cm;">{evidencia["titulo"]}</h3>')
                    
                    # Intentar mostrar la imagen (5x5 cm)
                    try:
                        url = evidencia['url']
                        # Si es una ruta local, construir la URL completa para el navegador
                        if not url.startswith('http'):
                            if url.startswith('/'):
                                url = f"{settings.MEDIA_URL.rstrip('/')}{url}"
                            else:
                                url = f"{settings.MEDIA_URL.rstrip('/')}/{url}"
                        
                        html_parts.append(f'<div style="text-align: center; margin: 0.5cm 0;">')
                        html_parts.append(f'<img src="{url}" alt="{evidencia.get("nombre", "")}" style="width: 5cm; height: 5cm; object-fit: contain; border: 1px solid #ddd;">')
                        html_parts.append('</div>')
                    except Exception as e:
                        # Si no se puede cargar la imagen, mostrar el nombre del archivo
                        html_parts.append(f'<p>Imagen: {evidencia.get("nombre", "No disponible")}</p>')
                    
                    html_parts.append('<div style="page-break-after: always;"></div>')  # Salto de página entre imágenes
                
                html_parts.append('</div>')
            
            # Agregar evidencias no-imagen en una sección separada
            evidencias_no_imagen = []
            for avance in evento_data['avances']:
                if avance.get('evidencias'):
                    for evidencia in avance['evidencias']:
                        if not evidencia.get('es_imagen'):
                            evidencias_no_imagen.append({
                                'avance': avance.get('descripcion_cambio', 'Sin título'),
                                'nombre': evidencia.get('nombre', ''),
                                'url': evidencia.get('url', '')
                            })
            
            if evidencias_no_imagen:
                if evidencias_imagenes:
                    html_parts.append('<div class="section-break">')
                html_parts.append('<h2>Evidencias (Archivos)</h2>')
                
                for evidencia in evidencias_no_imagen:
                    html_parts.append(f'<p><strong>Avance:</strong> {evidencia["avance"]}</p>')
                    html_parts.append(f'<p><strong>Archivo:</strong> {evidencia["nombre"]}</p>')
                    html_parts.append('<p></p>')
                
                if evidencias_imagenes:
                    html_parts.append('</div>')
        else:
            html_parts.append('<p>Este evento no tiene avances registrados.</p>')
    
    return '\n'.join(html_parts)


def generate_html_evento_individual(data):
    """Genera HTML para evento individual - Estructura completa con secciones"""
    print(f'🔵 generate_html_evento_individual llamado con data: {type(data)}')
    print(f'🔵 Keys en data: {list(data.keys()) if isinstance(data, dict) else "No es dict"}')
    
    html_parts = []
    
    if not data:
        return '<p>No se encontraron datos para este reporte.</p>'
    
    evento = data.get('evento', data)  # Puede venir como data.evento o directamente como data
    print(f'🔵 Evento extraído. Keys: {list(evento.keys()) if isinstance(evento, dict) else "No es dict"}')
    
    # ========== SECCIÓN 1: TÍTULO, FOTO DE PORTADA Y DATOS GENERALES EN LA MISMA HOJA ==========
    # Título del reporte
    html_parts.append('<h1 style="text-align: center; color: #0772d2; margin-bottom: 0.3cm;">Reporte de Evento Individual</h1>')
    html_parts.append(f'<p style="text-align: right; font-size: 9pt; color: #666; margin-bottom: 0.5cm;">Fecha de generación: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}</p>')
    
    # FOTO DE PORTADA (en la misma hoja)
    if evento.get('portada') and evento['portada'].get('url'):
        portada_url = evento['portada']['url']
        # Construir URL completa si es relativa
        if not portada_url.startswith('http'):
            if portada_url.startswith('/'):
                portada_url = f"{settings.MEDIA_URL.rstrip('/')}{portada_url}"
            else:
                portada_url = f"{settings.MEDIA_URL.rstrip('/')}/{portada_url}"
        
        html_parts.append('<div style="text-align: center; margin-bottom: 1cm;">')
        html_parts.append(f'<img src="{portada_url}" alt="Portada del evento" style="width: 100%; max-width: 15cm; height: 7.5cm; object-fit: cover; border-radius: 4px;">')
        html_parts.append('</div>')
    
    # DATOS GENERALES Y DESCRIPCIÓN (en la misma hoja, sin section-break)
    html_parts.append('<h2>Información General del Proyecto</h2>')
    html_parts.append('<div class="metric-box">')
    html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Nombre:</span> <strong style="color: #0772d2;">{evento.get("nombre", "-")}</strong></p>')
    html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Tipo:</span> <strong style="color: #0772d2;">{evento.get("tipo", "-")}</strong></p>')
    html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Estado:</span> <strong>{evento.get("estado", "-")}</strong></p>')
    html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Fecha:</span> <strong>{format_date(evento.get("fecha", evento.get("fecha_display", "-")))}</strong></p>')
    html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Ubicación:</span> <strong>{evento.get("ubicacion", evento.get("comunidad", "-"))}</strong></p>')
    html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Responsable:</span> <strong>{evento.get("responsable", "-")}</strong></p>')
    
    # Descripción
    if evento.get('descripcion'):
        html_parts.append('<p style="margin-top: 0.5cm; margin-bottom: 0.25cm;"><strong>Descripción:</strong></p>')
        html_parts.append(f'<p style="margin: 0.25cm 0; line-height: 1.6;">{evento.get("descripcion", "-")}</p>')
    
    html_parts.append('</div>')
    html_parts.append('</div>')
    
    # ========== SECCIÓN 3: PERSONAL ASIGNADO Y BENEFICIARIOS ==========
    html_parts.append('<div class="section-break">')
    
    # Personal asignado
    if evento.get('personal') and len(evento['personal']) > 0:
        html_parts.append('<h2>Personal Asignado</h2>')
        for persona in evento['personal']:
            html_parts.append('<p style="margin: 0.25cm 0;">')
            html_parts.append(f'• {persona.get("nombre", persona.get("username", "-"))}')
            if persona.get('puesto'):
                html_parts.append(f' - {persona.get("puesto", "")}')
            html_parts.append('</p>')
    else:
        html_parts.append('<h2>Personal Asignado</h2>')
        html_parts.append('<p>No hay personal asignado a este evento.</p>')
    
    html_parts.append('<p></p>')  # Espacio
    
    # Tabla de beneficiarios (en la misma hoja)
    if evento.get('beneficiarios') and len(evento['beneficiarios']) > 0:
        html_parts.append('<h2>Beneficiarios</h2>')
        html_parts.append('<table>')
        html_parts.append('<thead><tr><th>Nombre</th><th>Tipo</th><th>Comunidad</th><th>Información Adicional</th></tr></thead>')
        html_parts.append('<tbody>')
        
        for beneficiario in evento['beneficiarios']:
            html_parts.append('<tr>')
            html_parts.append(f'<td>{beneficiario.get("nombre", "-")}</td>')
            html_parts.append(f'<td>{beneficiario.get("tipo_display", beneficiario.get("tipo", "-"))}</td>')
            # Obtener comunidad de detalles
            comunidad_benef = '-'
            if beneficiario.get('detalles') and isinstance(beneficiario['detalles'], dict):
                comunidad_benef = beneficiario['detalles'].get('comunidad_nombre', '-')
            html_parts.append(f'<td>{comunidad_benef}</td>')
            html_parts.append(f'<td>{beneficiario.get("info_adicional", "-")}</td>')
            html_parts.append('</tr>')
        
        html_parts.append('</tbody></table>')
    else:
        html_parts.append('<h2>Beneficiarios</h2>')
        html_parts.append('<p>No hay beneficiarios registrados para este evento.</p>')
    
    html_parts.append('</div>')
    
    # ========== SECCIÓN 4: CAMBIOS/AVANCES (SIN EVIDENCIAS) ==========
    if evento.get('cambios') and len(evento['cambios']) > 0:
        html_parts.append('<div class="section-break">')
        html_parts.append('<h2>Cambios y Avances Realizados</h2>')
        html_parts.append('<table>')
        html_parts.append('<thead><tr><th>Descripción</th><th>Responsable</th><th>Comunidad donde se realizó el avance</th><th>Fecha y Hora</th></tr></thead>')
        html_parts.append('<tbody>')
        
        for cambio in evento['cambios']:
            html_parts.append('<tr>')
            html_parts.append(f'<td>{cambio.get("descripcion", "-")}</td>')
            html_parts.append(f'<td>{cambio.get("responsable", cambio.get("responsables_display", "-"))}</td>')
            # Obtener comunidades del cambio
            comunidades_cambio = cambio.get('comunidades', '')
            if not comunidades_cambio or comunidades_cambio == '':
                comunidades_cambio = 'No hay dato disponible'
            html_parts.append(f'<td>{comunidades_cambio}</td>')
            html_parts.append(f'<td>{cambio.get("fecha_display", cambio.get("fecha_cambio", "-"))}</td>')
            html_parts.append('</tr>')
        
        html_parts.append('</tbody></table>')
        html_parts.append('</div>')
    
    # ========== SECCIÓN 5: EVIDENCIAS (SOLO FOTOS) ==========
    # Recolectar todas las evidencias de imágenes de los cambios
    evidencias_imagenes = []
    if evento.get('cambios'):
        for cambio in evento['cambios']:
            if cambio.get('evidencias'):
                for evidencia in cambio['evidencias']:
                    # Solo agregar si es imagen
                    if evidencia.get('es_imagen') or (evidencia.get('tipo') and evidencia.get('tipo', '').startswith('image/')):
                        evidencias_imagenes.append({
                            'url': evidencia.get('url', ''),
                            'nombre': evidencia.get('nombre', ''),
                            'descripcion': evidencia.get('descripcion', ''),
                            'cambio_descripcion': cambio.get('descripcion', 'Sin descripción')
                        })
    
    if evidencias_imagenes:
        html_parts.append('<div class="section-break">')
        html_parts.append('<h2>Evidencias (Fotografías)</h2>')
        
        # Agrupar evidencias: hasta 2 por hoja
        for i in range(0, len(evidencias_imagenes), 2):
            # Si no es la primera iteración, agregar salto de página
            if i > 0:
                html_parts.append('<div class="section-break"></div>')
            
            # Primera evidencia
            if i < len(evidencias_imagenes):
                evidencia1 = evidencias_imagenes[i]
                
                # Título del avance al que pertenece
                html_parts.append(f'<h3 style="text-align: center; color: #0772d2; margin-top: 0.5cm; margin-bottom: 0.3cm;">{evidencia1["cambio_descripcion"]}</h3>')
                
                # Intentar mostrar la imagen (5x5 cm)
                url1 = evidencia1['url']
                if not url1.startswith('http'):
                    if url1.startswith('/'):
                        url1 = f"{settings.MEDIA_URL.rstrip('/')}{url1}"
                    else:
                        url1 = f"{settings.MEDIA_URL.rstrip('/')}/{url1}"
                
                html_parts.append(f'<div style="text-align: center; margin: 0.5cm 0;">')
                html_parts.append(f'<img src="{url1}" alt="{evidencia1.get("nombre", "")}" style="width: 5cm; height: 5cm; object-fit: contain; border: 1px solid #ddd;">')
                html_parts.append('</div>')
                html_parts.append('<p></p>')  # Espacio después de la imagen
            
            # Segunda evidencia (si existe)
            if i + 1 < len(evidencias_imagenes):
                evidencia2 = evidencias_imagenes[i + 1]
                
                # Título del avance al que pertenece
                html_parts.append(f'<h3 style="text-align: center; color: #0772d2; margin-top: 0.5cm; margin-bottom: 0.3cm;">{evidencia2["cambio_descripcion"]}</h3>')
                
                # Intentar mostrar la imagen (5x5 cm)
                url2 = evidencia2['url']
                if not url2.startswith('http'):
                    if url2.startswith('/'):
                        url2 = f"{settings.MEDIA_URL.rstrip('/')}{url2}"
                    else:
                        url2 = f"{settings.MEDIA_URL.rstrip('/')}/{url2}"
                
                html_parts.append(f'<div style="text-align: center; margin: 0.5cm 0;">')
                html_parts.append(f'<img src="{url2}" alt="{evidencia2.get("nombre", "")}" style="width: 5cm; height: 5cm; object-fit: contain; border: 1px solid #ddd;">')
                html_parts.append('</div>')
                html_parts.append('<p></p>')  # Espacio después de la imagen
        
        html_parts.append('</div>')
    
    # ========== SECCIÓN 6: GALERÍA DE IMÁGENES (ANEXOS) ==========
    # Obtener imágenes de la galería del evento (eventos_galeria)
    galeria_imagenes = []
    if evento.get('evidencias'):  # evidencias contiene las imágenes de la galería
        for imagen in evento['evidencias']:
            # Solo agregar si es imagen
            if imagen.get('es_imagen') or (imagen.get('tipo') and imagen.get('tipo', '').startswith('image/')):
                galeria_imagenes.append({
                    'url': imagen.get('url', ''),
                    'nombre': imagen.get('nombre', ''),
                    'descripcion': imagen.get('descripcion', '')
                })
    
    if galeria_imagenes:
        html_parts.append('<div class="section-break">')
        html_parts.append('<h2>Anexos - Galería de Imágenes</h2>')
        html_parts.append('<div style="display: flex; flex-wrap: wrap; justify-content: center; gap: 1cm;">')
        
        for imagen in galeria_imagenes:
            url = imagen['url']
            if not url.startswith('http'):
                if url.startswith('/'):
                    url = f"{settings.MEDIA_URL.rstrip('/')}{url}"
                else:
                    url = f"{settings.MEDIA_URL.rstrip('/')}/{url}"
            
            html_parts.append('<div style="text-align: center;">')
            html_parts.append(f'<img src="{url}" alt="{imagen.get("nombre", "")}" style="width: 5cm; height: 5cm; object-fit: contain; border: 1px solid #ddd; margin: 0.2cm;">')
            if imagen.get('descripcion'):
                html_parts.append(f'<p style="font-size: 8pt; color: #666; margin-top: 0.2cm;">{imagen.get("descripcion", "")}</p>')
            html_parts.append('</div>')
        
        html_parts.append('</div>')
        html_parts.append('</div>')
    
    return '\n'.join(html_parts)


def generate_html_comunidades(data):
    """Genera HTML para reporte de comunidades"""
    html_parts = []
    
    # Manejar caso donde data viene como dict con 'comunidades'
    if isinstance(data, dict) and 'comunidades' in data:
        data = data['comunidades']
    
    if not data or (isinstance(data, list) and len(data) == 0):
        return '<p>No se encontraron datos para este reporte.</p>'
    
    # Asegurar que data es una lista
    if not isinstance(data, list):
        return '<p>Error: formato de datos inválido para el reporte de comunidades.</p>'
    
    for idx, item in enumerate(data):
        # Validar que item es un diccionario
        if not isinstance(item, dict):
            continue
        
        # Salto de página antes de cada comunidad (excepto la primera)
        class_attr = ' class="section-break"' if idx > 0 else ''
        
        # ========== SECCIÓN 1: INFORMACIÓN GENERAL DE LA COMUNIDAD ==========
        # Título de la comunidad
        html_parts.append(f'<h2{class_attr}>{item.get("nombre", "N/A")}</h2>')
        
        # Información general con formato mejorado (todos en la misma hoja)
        html_parts.append('<div class="metric-box">')
        # COCODE
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">COCODE:</span> <strong style="color: #0772d2;">{item.get("cocode", "-")}</strong></p>')
        # Región
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Región:</span> <strong style="color: #0772d2;">{item.get("region", "-")}</strong></p>')
        # Tipo
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Tipo:</span> <strong style="color: #0772d2;">{item.get("tipo", "-")}</strong></p>')
        # Número de Beneficiarios
        total_beneficiarios = item.get('numero_beneficiarios', item.get('total_beneficiarios', 0))
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Número de Beneficiarios:</span> <strong style="color: #0772d2;">{total_beneficiarios}</strong></p>')
        # Número de Proyectos
        total_proyectos = item.get('numero_proyectos', item.get('total_proyectos', 0))
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Número de Proyectos:</span> <strong style="color: #0772d2;">{total_proyectos}</strong></p>')
        html_parts.append('</div>')
        
        # ========== SECCIÓN 2: TABLA DE PROYECTOS/EVENTOS ==========
        if item.get('proyectos') and len(item['proyectos']) > 0:
            html_parts.append('<h3>Proyectos/Eventos</h3>')
            html_parts.append('<table style="width: 100%; border-collapse: collapse; margin: 0.5cm 0; border: 1px solid #000;">')
            html_parts.append('<thead><tr>')
            html_parts.append('<th style="border: 1px solid #000; padding: 8px; background-color: #f0f0f0; text-align: left;">Nombre</th>')
            html_parts.append('<th style="border: 1px solid #000; padding: 8px; background-color: #f0f0f0; text-align: left;">Tipo</th>')
            html_parts.append('<th style="border: 1px solid #000; padding: 8px; background-color: #f0f0f0; text-align: left;">Estado</th>')
            html_parts.append('<th style="border: 1px solid #000; padding: 8px; background-color: #f0f0f0; text-align: left;">Fecha</th>')
            html_parts.append('</tr></thead>')
            html_parts.append('<tbody>')
            
            for proyecto in item['proyectos']:
                if isinstance(proyecto, dict):
                    html_parts.append('<tr>')
                    html_parts.append(f'<td style="border: 1px solid #000; padding: 8px;">{proyecto.get("nombre", "-")}</td>')
                    html_parts.append(f'<td style="border: 1px solid #000; padding: 8px;">{proyecto.get("tipo", "-")}</td>')
                    html_parts.append(f'<td style="border: 1px solid #000; padding: 8px;">{proyecto.get("estado", "-")}</td>')
                    html_parts.append(f'<td style="border: 1px solid #000; padding: 8px;">{format_date(proyecto.get("fecha", "-"))}</td>')
                    html_parts.append('</tr>')
            
            html_parts.append('</tbody>')
            html_parts.append('</table>')
    
    return '\n'.join(html_parts)


def generate_html_actividad_usuarios(data):
    """Genera HTML para actividad de usuarios"""
    html_parts = []
    
    # Manejar caso donde data viene como dict con 'usuarios'
    if isinstance(data, dict) and 'usuarios' in data:
        data = data['usuarios']
    
    if not data or (isinstance(data, list) and len(data) == 0):
        return '<p>No se encontraron datos para este reporte.</p>'
    
    # Asegurar que data es una lista
    if not isinstance(data, list):
        return '<p>Error: formato de datos inválido para el reporte de actividad de usuarios.</p>'
    
    for idx, item in enumerate(data):
        # Validar que item es un diccionario
        if not isinstance(item, dict):
            continue
        
        # Salto de página antes de cada usuario (excepto el primero)
        class_attr = ' class="section-break"' if idx > 0 else ''
        
        # ========== SECCIÓN: INFORMACIÓN DEL USUARIO ==========
        nombre_usuario = item.get('nombre', item.get('username', 'N/A'))
        html_parts.append(f'<h2{class_attr}>Usuario: {nombre_usuario}</h2>')
        
        # Información general del usuario
        html_parts.append('<div class="metric-box">')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Usuario:</span> <strong style="color: #0772d2;">{item.get("username", "-")}</strong></p>')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Rol:</span> <strong style="color: #0772d2;">{item.get("rol_display", item.get("rol", "-"))}</strong></p>')
        if item.get('puesto_nombre') or item.get('puesto'):
            puesto = item.get('puesto_nombre', item.get('puesto', '-'))
            html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Puesto:</span> <strong style="color: #0772d2;">{puesto}</strong></p>')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Total de cambios:</span> <strong style="color: #0772d2;">{item.get("total_cambios", 0)}</strong></p>')
        html_parts.append('</div>')
        
        # ========== SECCIÓN: TABLA DE CAMBIOS POR PROYECTO ==========
        if item.get('cambios') and len(item['cambios']) > 0:
            # Agrupar cambios por evento/proyecto
            cambios_por_evento = {}
            for cambio in item['cambios']:
                if not isinstance(cambio, dict):
                    continue
                
                evento_id = cambio.get('evento_id', 'sin_evento')
                evento_nombre = cambio.get('evento_nombre', 'Sin evento')
                
                if evento_id not in cambios_por_evento:
                    cambios_por_evento[evento_id] = {
                        'nombre': evento_nombre,
                        'cambios': []
                    }
                
                cambios_por_evento[evento_id]['cambios'].append(cambio)
            
            html_parts.append('<h3>Tabla de Cambios Realizados</h3>')
            html_parts.append('<table style="width: 100%; border-collapse: collapse; margin: 0.5cm 0; border: 1px solid #000;">')
            html_parts.append('<thead><tr>')
            html_parts.append('<th style="border: 1px solid #000; padding: 8px; background-color: #f0f0f0; text-align: left;">Proyecto</th>')
            html_parts.append('<th style="border: 1px solid #000; padding: 8px; background-color: #f0f0f0; text-align: left;">Tipo de Cambio</th>')
            html_parts.append('<th style="border: 1px solid #000; padding: 8px; background-color: #f0f0f0; text-align: left;">Descripción</th>')
            html_parts.append('<th style="border: 1px solid #000; padding: 8px; background-color: #f0f0f0; text-align: left;">Fecha y Hora</th>')
            html_parts.append('</tr></thead>')
            html_parts.append('<tbody>')
            
            # Agregar filas agrupadas por proyecto
            for evento_id, evento_data in cambios_por_evento.items():
                cambios_evento = evento_data['cambios']
                
                # Agregar fila con nombre del proyecto (primera fila del grupo)
                if len(cambios_evento) > 0:
                    html_parts.append(f'<tr style="background-color: #e6f2ff;">')
                    html_parts.append(f'<td style="border: 1px solid #000; padding: 8px; font-weight: bold; color: #0772d2;" colspan="4">{evento_data["nombre"]}</td>')
                    html_parts.append('</tr>')
                    
                    # Agregar filas de cambios para este proyecto
                    for cambio in cambios_evento:
                        html_parts.append('<tr>')
                        html_parts.append('<td style="border: 1px solid #000; padding: 8px;"></td>')  # Proyecto ya está en la fila anterior
                        html_parts.append(f'<td style="border: 1px solid #000; padding: 8px;">{cambio.get("tipo_cambio", "-")}</td>')
                        html_parts.append(f'<td style="border: 1px solid #000; padding: 8px;">{cambio.get("descripcion", "-")}</td>')
                        # Usar fecha_display si está disponible
                        fecha_str = cambio.get('fecha_display', '')
                        if not fecha_str and cambio.get('fecha'):
                            fecha_str = format_date(cambio.get('fecha', '-'))
                        html_parts.append(f'<td style="border: 1px solid #000; padding: 8px;">{fecha_str if fecha_str else "-"}</td>')
                        html_parts.append('</tr>')
            
            html_parts.append('</tbody>')
            html_parts.append('</table>')
    
    return '\n'.join(html_parts)


def generate_html_reporte_general(data):
    """Genera HTML para reporte general - Datos primero, luego tablas"""
    html_parts = []
    
    if not data:
        return '<p>No se encontraron datos para este reporte.</p>'
    
    # ========== SECCIÓN 1: TODOS LOS DATOS EN UNA MISMA HOJA ==========
    html_parts.append('<h2>Resumen General</h2>')
    
    # Tipos de eventos (dato)
    if 'tipos_eventos' in data:
        tipos = data['tipos_eventos']
        html_parts.append('<div class="metric-box">')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Total de Tipos de Evento:</span> <strong style="color: #0772d2;">Capacitación: {tipos.get("capacitacion", 0)}, Entrega: {tipos.get("entrega", 0)}, Proyecto de Ayuda: {tipos.get("proyecto_ayuda", 0)}</strong></p>')
        html_parts.append('</div>')
    
    # Total de beneficiarios alcanzado (dato)
    if 'total_beneficiarios' in data:
        html_parts.append('<div class="metric-box">')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Total de Beneficiarios Alcanzados:</span> <strong style="color: #0772d2; font-size: 14pt;">{data["total_beneficiarios"]}</strong></p>')
        html_parts.append('</div>')
    
    # Tipos de beneficiarios (dato)
    if 'tipos_beneficiarios' in data:
        tipos = data['tipos_beneficiarios']
        html_parts.append('<div class="metric-box">')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Total de Tipos de Beneficiarios:</span> <strong style="color: #0772d2;">Individuales: {tipos.get("individual", 0)}, Familias: {tipos.get("familia", 0)}, Instituciones: {tipos.get("institucion", 0)}, Mujeres: {tipos.get("mujeres", 0)}, Hombres: {tipos.get("hombres", 0)}</strong></p>')
        html_parts.append('</div>')
    
    # Total de comunidades alcanzadas (dato - solo el número)
    if 'total_comunidades' in data:
        total_comunidades = data['total_comunidades']
        html_parts.append('<div class="metric-box">')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Total de Comunidades Alcanzadas:</span> <strong style="color: #0772d2; font-size: 14pt;">{total_comunidades.get("total", 0)}</strong></p>')
        html_parts.append('</div>')
    
    # Total de avances en Proyectos (dato)
    if 'total_avances' in data:
        html_parts.append('<div class="metric-box">')
        html_parts.append(f'<p style="margin: 0.25cm 0;"><span class="metric-label">Total de Avances en Proyectos:</span> <strong style="color: #0772d2; font-size: 14pt;">{data["total_avances"]}</strong></p>')
        html_parts.append('</div>')
    
    # Evento con más avances y cambios (pequeña tabla)
    if 'evento_mas_avances' in data and data['evento_mas_avances']:
        evento = data['evento_mas_avances']
        html_parts.append('<h3>Evento con Más Avances y Cambios</h3>')
        html_parts.append('<table style="width: 80%; margin: 0 auto;">')
        html_parts.append('<thead><tr><th>Nombre</th><th>Total de Avances</th><th>Total de Cambios</th></tr></thead>')
        html_parts.append('<tbody>')
        html_parts.append('<tr>')
        html_parts.append(f'<td>{evento.get("nombre", "-")}</td>')
        html_parts.append(f'<td>{evento.get("total_avances", 0)}</td>')
        html_parts.append(f'<td>{evento.get("total_cambios", 0)}</td>')
        html_parts.append('</tr>')
        html_parts.append('</tbody></table>')
    
    # ========== SECCIÓN 2: TABLAS - CADA UNA EN SU PROPIA HOJA ==========
    
    # Tabla 1: Total de eventos
    if 'total_eventos' in data:
        total_eventos = data['total_eventos']
        if total_eventos.get('lista_eventos') and len(total_eventos['lista_eventos']) > 0:
            html_parts.append('<div class="section-break"><h2>Total de Eventos</h2>')
            html_parts.append(f'<div class="metric-box"><p style="margin: 0.25cm 0;"><span class="metric-label">Total:</span> <strong>{total_eventos.get("total", 0)}</strong></p></div>')
            html_parts.append('<table>')
            html_parts.append('<thead><tr><th>Nombre</th><th>Tipo</th><th>Estado</th><th>Fecha</th><th>Comunidad</th><th>Beneficiarios</th></tr></thead>')
            html_parts.append('<tbody>')
            
            for evento in total_eventos['lista_eventos']:
                html_parts.append('<tr>')
                html_parts.append(f'<td>{evento.get("nombre", "-")}</td>')
                html_parts.append(f'<td>{evento.get("tipo", "-")}</td>')
                html_parts.append(f'<td>{evento.get("estado", "-")}</td>')
                html_parts.append(f'<td>{format_date(evento.get("fecha", "-"))}</td>')
                html_parts.append(f'<td>{evento.get("comunidad", "-")}</td>')
                html_parts.append(f'<td>{evento.get("beneficiarios", 0)}</td>')
                html_parts.append('</tr>')
            
            html_parts.append('</tbody></table></div>')
    
    # Tabla 2: Total de comunidades alcanzadas (si tiene lista)
    if 'total_comunidades' in data:
        total_comunidades = data['total_comunidades']
        if total_comunidades.get('lista_comunidades') and len(total_comunidades['lista_comunidades']) > 0:
            html_parts.append('<div class="section-break"><h2>Total de Comunidades Alcanzadas</h2>')
            html_parts.append(f'<div class="metric-box"><p style="margin: 0.25cm 0;"><span class="metric-label">Total:</span> <strong>{total_comunidades.get("total", 0)}</strong></p></div>')
            html_parts.append('<table>')
            html_parts.append('<thead><tr><th>Comunidad</th><th>Región</th></tr></thead>')
            html_parts.append('<tbody>')
            
            for comunidad in total_comunidades['lista_comunidades']:
                html_parts.append('<tr>')
                html_parts.append(f'<td>{comunidad.get("nombre", "-")}</td>')
                html_parts.append(f'<td>{comunidad.get("region", "-")}</td>')
                html_parts.append('</tr>')
            
            html_parts.append('</tbody></table></div>')
    
    # Tabla 3: Comunidades con más beneficiarios y eventos
    if 'comunidades_mas_beneficiarios' in data and len(data['comunidades_mas_beneficiarios']) > 0:
        html_parts.append('<div class="section-break"><h2>Comunidades con Más Beneficiarios y Eventos</h2>')
        comunidades = data['comunidades_mas_beneficiarios']
        html_parts.append('<table>')
        html_parts.append('<thead><tr><th>Comunidad</th><th>Total de Beneficiarios</th><th>Total de Eventos</th></tr></thead>')
        html_parts.append('<tbody>')
        
        for com in comunidades:
            html_parts.append('<tr>')
            html_parts.append(f'<td>{com.get("nombre", "-")}</td>')
            html_parts.append(f'<td>{com.get("total_beneficiarios", 0)}</td>')
            html_parts.append(f'<td>{com.get("total_eventos", 0)}</td>')
            html_parts.append('</tr>')
        
        html_parts.append('</tbody></table></div>')
    
    # Tabla 4: Eventos con más beneficiarios
    if 'eventos_mas_beneficiarios' in data and len(data['eventos_mas_beneficiarios']) > 0:
        html_parts.append('<div class="section-break"><h2>Eventos con Más Beneficiarios</h2>')
        eventos = data['eventos_mas_beneficiarios']
        html_parts.append('<table>')
        html_parts.append('<thead><tr><th>Evento</th><th>Total de Beneficiarios</th></tr></thead>')
        html_parts.append('<tbody>')
        
        for evento in eventos:
            html_parts.append('<tr>')
            html_parts.append(f'<td>{evento.get("nombre", "-")}</td>')
            html_parts.append(f'<td>{evento.get("total_beneficiarios", 0)}</td>')
            html_parts.append('</tr>')
        
        html_parts.append('</tbody></table></div>')
    
    return '\n'.join(html_parts)

